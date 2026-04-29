#!/usr/bin/env python3
"""
Organizador de Documentos Juridicos - AB Group
App web mobile-first para organizar documentos em ordem cronologica via IA.
Separa automaticamente PDFs com multiplos documentos em arquivos individuais.
"""

import os
import io
import re
import sys
import json
import uuid
import time
import base64
import shutil
import logging
import sqlite3
import zipfile
import tempfile
import unicodedata
import threading
from pathlib import Path
from datetime import datetime, timedelta

# Logging persistente em stdout — Render captura e mantem por 7 dias.
# Pode ser auditado via Render Dashboard > Logs sempre que precisar.
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)],
)
audit_log = logging.getLogger("audit")

# ============ SQLite local pra auditoria persistente ============
# /tmp pode ser zerado em deploys, mas sobrevive a restarts do worker.
# Para persistencia real entre deploys, migrar pra VPS (#4.2) ou usar
# servico externo (Sentry, Logflare, etc).
AUDIT_DB_PATH = Path(tempfile.gettempdir()) / "organizador-audit.db"
_audit_lock = threading.Lock()


def _audit_init_db():
    """Cria tabela de auditoria se nao existir."""
    try:
        with _audit_lock:
            conn = sqlite3.connect(str(AUDIT_DB_PATH))
            conn.execute("""
                CREATE TABLE IF NOT EXISTS audit_log (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    timestamp TEXT NOT NULL,
                    evento TEXT NOT NULL,           -- 'INICIO' | 'SUCESSO' | 'ERRO'
                    usuario TEXT,
                    cliente TEXT,
                    tipo_processo TEXT,
                    arquivos TEXT,                  -- nomes (csv)
                    total_documentos INTEGER,
                    tipo_erro TEXT,
                    mensagem TEXT,
                    detalhes TEXT                   -- JSON livre
                )
            """)
            conn.execute("CREATE INDEX IF NOT EXISTS idx_timestamp ON audit_log(timestamp)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_evento ON audit_log(evento)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_usuario ON audit_log(usuario)")
            conn.commit()
            conn.close()
    except Exception as e:
        # Falha na inicializacao do DB nao deve quebrar o app
        print(f"AUDIT_DB_INIT_ERROR: {e}", file=sys.stderr)


def audit_save(evento, usuario=None, cliente=None, tipo_processo=None,
              arquivos=None, total_documentos=0, tipo_erro=None, mensagem=None,
              detalhes=None):
    """Salva uma linha de auditoria no SQLite local + stdout."""
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    arq_str = ""
    if arquivos:
        if isinstance(arquivos, (list, tuple)):
            arq_str = ",".join(str(a)[:50] for a in arquivos)[:1000]
        else:
            arq_str = str(arquivos)[:1000]
    msg = (str(mensagem)[:500] if mensagem else "")
    det = json.dumps(detalhes, ensure_ascii=False)[:2000] if detalhes else None

    try:
        with _audit_lock:
            conn = sqlite3.connect(str(AUDIT_DB_PATH))
            conn.execute(
                "INSERT INTO audit_log(timestamp, evento, usuario, cliente, tipo_processo, "
                "arquivos, total_documentos, tipo_erro, mensagem, detalhes) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (ts, evento, usuario, cliente, tipo_processo, arq_str,
                 int(total_documentos or 0), tipo_erro, msg, det),
            )
            conn.commit()
            conn.close()
    except Exception as e:
        print(f"AUDIT_DB_INSERT_ERROR: {e}", file=sys.stderr)


def audit_query(dias=7, limit=500):
    """Le ultimos N dias do SQLite. Retorna lista de dicts."""
    try:
        with _audit_lock:
            conn = sqlite3.connect(str(AUDIT_DB_PATH))
            conn.row_factory = sqlite3.Row
            cutoff = (datetime.now() - timedelta(days=dias)).strftime("%Y-%m-%d %H:%M:%S")
            rows = conn.execute(
                "SELECT * FROM audit_log WHERE timestamp >= ? ORDER BY timestamp DESC LIMIT ?",
                (cutoff, limit),
            ).fetchall()
            conn.close()
            return [dict(r) for r in rows]
    except Exception as e:
        print(f"AUDIT_DB_READ_ERROR: {e}", file=sys.stderr)
        return []


# Inicializa DB no startup
_audit_init_db()

from flask import Flask, render_template, request, jsonify, send_file, Response, stream_with_context

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024  # 100MB max total

# Diretorio compartilhado entre workers para ZIPs temporarios
SHARED_ZIP_DIR = Path(tempfile.gettempdir()) / "organizador-zips"
SHARED_ZIP_DIR.mkdir(exist_ok=True)

# Diretorio para sessoes de revisao manual (#2.2)
# Cada sessao contem os arquivos processados + metadata.json
# Permite usuario revisar/editar a classificacao antes de gerar o ZIP final
SESSIONS_DIR = Path(tempfile.gettempdir()) / "organizador-sessions"
SESSIONS_DIR.mkdir(exist_ok=True)

MODELO_IA = "claude-haiku-4-5-20251001"
EXTENSOES_ACEITAS = {".pdf", ".jpg", ".jpeg", ".png", ".docx"}
MAX_TEXTO_POR_PAGINA = 2000  # contexto maior para IA classificar melhor
MAX_PAGINAS_PDF = 100  # suporta PDFs grandes (antes era 30 = perdia paginas)

# Limite de tamanho por tipo de processo (em bytes)
LIMITES_TAMANHO = {
    "inss_admin": 5 * 1024 * 1024,   # 5MB
    "judicial": 10 * 1024 * 1024,    # 10MB
    "consumidor": 10 * 1024 * 1024,  # 10MB
    "trabalhista": 10 * 1024 * 1024, # 10MB
    "civel": 10 * 1024 * 1024,       # 10MB
}
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # fallback geral

TIPOS_PROCESSO = {
    "inss_admin": "INSS Administrativo",
    "judicial": "Judicial",
    "consumidor": "Direito do Consumidor",
    "trabalhista": "Trabalhista",
    "civel": "Cível",
}

USUARIOS = [
    "Emily", "Karol", "Alan", "Henrique", "Caique", "Jaíne",
    "Camila", "Luana", "Claudio", "Meire", "André", "Vitória",
]

# Webhook para registrar uso e erros (mesma planilha, diferenciado pelo prefixo no status)
LOG_WEBHOOK_URL = os.environ.get("LOG_WEBHOOK_URL", "")

# Dashboard admin (#5.4) — Andre + Dra. Luana
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "abgroup2026")  # default p/ dev — sobrescrever em prod
SPREADSHEET_ID = "182DGXAm25l1AVfgdA6uDxU4BvpWAxA7GxWQLKRlsVIc"  # planilha de logs
# URL publica da planilha em CSV (requer planilha "qualquer pessoa com link pode ver")
SPREADSHEET_CSV_URL = f"https://docs.google.com/spreadsheets/d/{SPREADSHEET_ID}/gviz/tq?tqx=out:csv"

PROMPT_MAPEAMENTO = """Analise o texto de cada pagina deste PDF juridico brasileiro.
Identifique TODOS os documentos separados que existem dentro deste arquivo.

IMPORTANTE: Responda SEMPRE em portugues brasileiro. NUNCA use ingles.
Use EXATAMENTE um destes tipos (escreva identico, sem traduzir):
- Procuracao
- Substabelecimento
- Declaracao de Hipossuficiencia     (cliente declara que nao pode pagar custas)
- Declaracao de Residencia            (cliente declara onde mora)
- Declaracao de Beneficios INSS       (documento DO INSS listando beneficios)
- Declaracao                          (outras declaracoes do cliente)
- Contrato de Honorarios
- Termo de Responsabilidade
- Termo de Representacao
- Protocolo de Assinatura             (gerado por DocuSign, Clicksign, ZapSign, D4Sign, etc)
- Relatorio de Assinaturas            (relatorio ZapSign com assinantes e tokens)
- Comunicacao de Decisao              (documento DO INSS com decisao sobre beneficio)
- Carta Indeferimento INSS
- Peticao Inicial                     (peticao do advogado iniciando o processo judicial)
- Calculo do Valor da Causa           (planilha/documento com calculo do valor pedido)
- RG
- CPF
- CNH
- Documento de Identidade
- CNIS
- CTPS
- Atestado Medico
- Relatorio Medico
- Receita Medica
- Exame Medico
- Laudo Medico
- Laudo MeuINSS
- Certidao de Casamento
- Certidao de Nascimento
- Certidao de Obito
- Termo de Homologacao Atividade Rural
- Despacho Decisorio                  (decisao administrativa que considerou periodo de Segurado Especial)
- Documentos Rurais
- Folha V7
- Declaracao Tempo de Servico
- Ficha Funcionario
- Certidao Tempo de Servico
- Ficha Financeira
- PPP
- LTCAT
- GPS
- Comprovante de Residencia
- Comprovante de Gasto
- Foto de Residencia
- Avaliacao Social
- Pericia Medica
- Quadro de Informacoes da Avaliacao Social e Pericia Medica
- Contagem de Tempo
- Calculo Regras Transicao
- Calculo da Renda Mensal Inicial     (RMI)
- Copia Processo Administrativo
- Certidao Negativa Justica Estadual
- Protocolo de Requerimento INSS      (formulario de pedido de beneficio no MeuINSS)
- Print TRF / Print Processo          (lista de processos judiciais)

REGRAS CRITICAS DE DIFERENCIACAO:
1. "Declaracao de Beneficios" (do INSS) NAO e a mesma coisa que "Declaracao de Hipossuficiencia"
   (do cliente). O INSS emite a de Beneficios. O cliente emite a de Hipossuficiencia.
2. "Comunicacao de Decisao" (INSS informa negado/concedido) NAO e "Declaracao".
3. "Termo de Representacao" deve ser classificado como "Termo de Responsabilidade" na saida
   (e o mesmo tipo pro fluxo).
4. "Relatorio de Assinaturas" da ZapSign/Clicksign deve ser classificado como
   "Protocolo de Assinatura" (e o mesmo tipo).

REGRAS GERAIS:
- Cada documento pode ter 1 ou mais paginas
- Uma procuracao de 3 paginas NAO e 3 procuracoes - e UMA procuracao (pagina_inicio=1, pagina_fim=3)
- "Protocolo de Assinatura" vem APOS o documento assinado
- A data deve ser a data de emissao/expedicao (no texto), NAO a data de hoje

Responda APENAS em JSON valido, sem markdown:
{"documentos": [
  {"tipo": "Procuracao", "pagina_inicio": 1, "pagina_fim": 1, "data": "2026-03-14"},
  {"tipo": "Protocolo de Assinatura", "pagina_inicio": 2, "pagina_fim": 2, "data": null},
  {"tipo": "Declaracao de Residencia", "pagina_inicio": 3, "pagina_fim": 3, "data": "2026-03-14"},
  {"tipo": "Declaracao de Beneficios INSS", "pagina_inicio": 4, "pagina_fim": 4, "data": "2025-10-17"},
  {"tipo": "Comunicacao de Decisao", "pagina_inicio": 5, "pagina_fim": 5, "data": "2026-04-02"}
]}

Se nao encontrar data, use "data": null.
"""

PROMPT_EXTRACAO_SIMPLES = """Analise este documento juridico brasileiro.

IMPORTANTE: Responda SEMPRE em portugues brasileiro, nunca em ingles.

Extraia:
1. Tipo do documento. Use EXATAMENTE um destes (copiar identico):
   Peticao Inicial, Calculo do Valor da Causa, Procuracao, Substabelecimento,
   Declaracao de Hipossuficiencia, Declaracao de Residencia, Declaracao de Beneficios INSS,
   Declaracao, Contrato de Honorarios, Termo de Responsabilidade, Termo de Representacao,
   Protocolo de Assinatura, Relatorio de Assinaturas, Comunicacao de Decisao,
   Carta Indeferimento INSS, RG, CPF, CNH, CNIS, CTPS, Atestado Medico, Relatorio Medico,
   Receita Medica, Exame Medico, Laudo Medico, Laudo MeuINSS, Certidao de Casamento,
   Certidao de Nascimento, Certidao de Obito, Documentos Rurais, Folha V7,
   Termo de Homologacao Atividade Rural, Despacho Decisorio,
   Declaracao Tempo Servico, Ficha Funcionario, Certidao Tempo Servico, Ficha Financeira,
   PPP, LTCAT, GPS, Comprovante Residencia, Comprovante Gasto, Foto Residencia,
   Protocolo de Requerimento INSS, Print TRF, Pericia Medica,
   Quadro de Informacoes Avaliacao Social, Contagem de Tempo, Calculo Regras Transicao,
   Calculo RMI, Copia Processo Administrativo, Certidao Negativa Justica Estadual
2. Data de emissao/expedicao DO DOCUMENTO (procure no TOPO, RODAPE ou DATA EXPLICITA no texto).
   Formatos possiveis: DD/MM/YYYY, DD-MM-YYYY, "14 de Marco de 2026", YYYY-MM-DD.
   NAO use a data atual nem data de impressao do sistema.

DIFERENCIACAO IMPORTANTE:
- "Declaracao de Beneficios" e um DOCUMENTO DO INSS listando beneficios. NUNCA e do cliente.
- "Declaracao de Hipossuficiencia" e do CLIENTE declarando ser pobre pra nao pagar custas.
- "Declaracao de Residencia" e do CLIENTE declarando onde mora.
- "Comunicacao de Decisao" e do INSS informando negado/concedido.

Responda APENAS em JSON valido, sem markdown:
{"tipo": "...", "data": "YYYY-MM-DD"}

Se nao encontrar data, use "data": null.
"""

PROMPT_MAPEAMENTO_CURTO = """Identifique os documentos nestas paginas de PDF juridico brasileiro.

IMPORTANTE: Responda SEMPRE em portugues brasileiro. Uma procuracao de 3 paginas e UMA procuracao (nao 3).
Tipos permitidos: Peticao Inicial, Calculo do Valor da Causa, Procuracao, Substabelecimento,
Declaracao de Hipossuficiencia, Declaracao, Contrato de Honorarios, Termo de Responsabilidade,
Termo de Representacao, Protocolo de Assinatura, Comunicacao de Decisao, Carta Indeferimento INSS,
RG, CPF, CNH, CNIS, CTPS, Atestado Medico, Relatorio Medico, Receita Medica, Exame Medico,
Laudo, Certidao, Despacho Decisorio, PPP, LTCAT, GPS, Comprovante Residencia, Comprovante Gasto,
Quadro de Informacoes Avaliacao Social, Contagem de Tempo, Calculo Regras Transicao, Calculo RMI,
Copia Processo Administrativo, Certidao Negativa Justica Estadual.

Para cada documento, indique: tipo, pagina de inicio, pagina de fim, data de emissao (do texto, nao hoje).

Responda APENAS em JSON valido:
{"documentos": [{"tipo": "...", "pagina_inicio": 1, "pagina_fim": 3, "data": "YYYY-MM-DD"}]}
Se nao encontrar data, use "data": null.
"""

CHUNK_SIZE_PAGINAS = 8  # paginas por chunk para PDFs grandes

# Categorias que passam por DUPLA CHECAGEM (#3.2): 2a chamada da IA
# pra confirmar a classificacao. Aplicada em tipos onde erro tem alto custo
# (procuracao, declaracoes, contratos, decisoes do INSS).
TIPOS_DUPLA_CHECAGEM = {
    "procuracao",
    "substabelecimento",
    "termo_de_responsabilidade",
    "protocolo_assinatura",
    "declaracao",
    "declaracao_hipossuficiencia",
    "declaracao_beneficios_inss",
    "contrato_de_honorarios",
    "carta_indeferimento",
    "cnis",
}

PROMPT_DUPLA_CHECAGEM = """Voce esta validando a classificacao de um documento juridico brasileiro.
Outra IA classificou este documento como: "{tipo_sugerido}"

Sua tarefa: olhar o conteudo e decidir se a classificacao esta CORRETA ou se deveria ser OUTRA.

Tipos validos (use EXATAMENTE um destes):
Procuracao, Substabelecimento, Termo de Representacao, Termo de Responsabilidade,
Protocolo de Assinatura, Relatorio de Assinaturas, Declaracao de Hipossuficiencia,
Declaracao de Residencia, Declaracao de Beneficios INSS, Declaracao,
Contrato de Honorarios, RG, CPF, CNH, CNIS, CTPS, Comunicacao de Decisao,
Carta Indeferimento INSS, Atestado Medico, Relatorio Medico, Receita Medica,
Exame Medico, Laudo Medico, Pericia Medica, Comprovante Residencia, Comprovante Gasto,
Foto Residencia, PPP, LTCAT, GPS, Print TRF, Protocolo de Requerimento INSS, desconhecido

REGRAS DE DESEMPATE:
- "Declaracao de Beneficios INSS" e DOCUMENTO DO INSS (vem do MeuINSS, lista beneficios)
- "Declaracao de Hipossuficiencia" e DO CLIENTE (declara pobreza)
- "Comunicacao de Decisao" e DO INSS (informa NEGADO/CONCEDIDO)
- "Carta Indeferimento INSS" e DO INSS (informa indeferimento)
- "Relatorio de Assinaturas" e do ZapSign/DocuSign (lista quem assinou)
- "Protocolo de Assinatura" tambem e do servico de assinatura digital

Responda APENAS em JSON valido, sem markdown:
{{"tipo_correto": "...", "concorda": true/false, "razao": "breve justificativa em 1 frase"}}
"""


def hash_arquivo_binario(caminho):
    """Hash SHA256 do binario completo. Detecta duplicatas EXATAS."""
    import hashlib
    h = hashlib.sha256()
    try:
        with open(caminho, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return None


def hash_conteudo_pdf(caminho):
    """Hash do TEXTO extraido (primeiras 3 paginas). Detecta duplicatas com
    pequenas diferencas de metadata/timestamps."""
    try:
        import hashlib
        import pdfplumber
        textos = []
        with pdfplumber.open(caminho) as pdf:
            for p in pdf.pages[:3]:
                t = p.extract_text() or ""
                # Normaliza: remove whitespace, lowercase
                t = " ".join(t.split()).lower()
                textos.append(t)
        texto_total = " ".join(textos)
        if len(texto_total) < 50:  # texto muito curto, nao confiavel
            return None
        return hashlib.sha256(texto_total.encode("utf-8")).hexdigest()
    except Exception:
        return None


def detectar_duplicatas(docs):
    """Marca duplicatas no doc com:
      - _duplicata: True se for duplicata
      - _duplicata_de: id (indice) do doc original
      - _motivo_duplicata: 'arquivo_identico' | 'mesmo_conteudo'
    A duplicata fica marcada pra ser ignorada (deletar=True por padrao na revisao).
    Retorna numero de duplicatas detectadas.
    """
    seen_bin = {}    # hash_bin -> indice do primeiro doc
    seen_txt = {}    # hash_txt -> indice do primeiro doc
    duplicatas = 0

    for i, doc in enumerate(docs):
        arq = doc.get("arquivo_tmp")
        if not arq or not os.path.exists(arq):
            continue

        h_bin = hash_arquivo_binario(arq)
        h_txt = None
        if Path(arq).suffix.lower() == ".pdf":
            h_txt = hash_conteudo_pdf(arq)

        # Hash binario igual = duplicata exata
        if h_bin and h_bin in seen_bin:
            doc["_duplicata"] = True
            doc["_duplicata_de"] = seen_bin[h_bin]
            doc["_motivo_duplicata"] = "arquivo_identico"
            duplicatas += 1
            continue

        # Hash texto igual = mesmo conteudo (com diferencas minimas)
        if h_txt and h_txt in seen_txt:
            doc["_duplicata"] = True
            doc["_duplicata_de"] = seen_txt[h_txt]
            doc["_motivo_duplicata"] = "mesmo_conteudo"
            duplicatas += 1
            continue

        # Nao e duplicata — registra
        if h_bin:
            seen_bin[h_bin] = i
        if h_txt:
            seen_txt[h_txt] = i

    return duplicatas


def eh_tipo_critico(tipo_str):
    """Verifica se um tipo precisa de dupla checagem."""
    fake_doc = {"tipo": tipo_str}
    cat = classificar_doc(fake_doc, "inss_admin")  # tipo_processo nao importa pra detectar critico
    return cat in TIPOS_DUPLA_CHECAGEM


def dupla_checagem_doc(client, doc, max_chars=2000):
    """Para tipos criticos, faz 2a chamada da IA pra confirmar classificacao.
    Adiciona 2 campos ao doc (in-place):
      - _confianca: "alta" se concordou, "media" se divergiu, None se nao foi checado
      - _tipo_alternativo: tipo sugerido pela 2a IA (so se divergiu)
    Mantem o tipo original no doc — usuario decide na tela de revisao se quer trocar.
    """
    tipo_atual = doc.get("tipo", "")
    if not eh_tipo_critico(tipo_atual):
        return  # nao e critico, nao precisa

    arquivo = doc.get("arquivo_tmp")
    if not arquivo or not os.path.exists(arquivo):
        return

    ext = Path(arquivo).suffix.lower()
    content_msg = []

    try:
        if ext == ".pdf":
            # Tenta texto da pagina 1; se vazio, usa imagem
            import pdfplumber
            with pdfplumber.open(arquivo) as pdf:
                if pdf.pages:
                    texto = (pdf.pages[0].extract_text() or "").strip()
                else:
                    texto = ""
            if texto:
                content_msg.append({"type": "text", "text": texto[:max_chars]})
            else:
                img_b64 = pagina_para_imagem_b64(arquivo, 1)
                if img_b64:
                    content_msg.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64},
                    })
        elif ext in (".jpg", ".jpeg", ".png"):
            with open(arquivo, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode()
            content_msg.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64},
            })
        elif ext == ".docx":
            from docx import Document
            d = Document(arquivo)
            texto = "\n".join(p.text for p in d.paragraphs if p.text.strip())[:max_chars]
            if texto:
                content_msg.append({"type": "text", "text": texto})
    except Exception:
        return  # se falhou em ler, deixa sem checagem

    if not content_msg:
        return

    prompt = PROMPT_DUPLA_CHECAGEM.format(tipo_sugerido=tipo_atual)
    content_msg.append({"type": "text", "text": prompt})

    try:
        response = client.messages.create(
            model=MODELO_IA,
            max_tokens=200,
            messages=[{"role": "user", "content": content_msg}],
        )
        dados = _parse_json_response(response.content[0].text)
        tipo_correto = (dados.get("tipo_correto") or "").strip()
        concorda = dados.get("concorda", True)

        if not tipo_correto:
            return  # resposta invalida — ignora
        # Compara case-insensitive (mas mantem o original)
        if concorda or tipo_correto.lower() == tipo_atual.lower():
            doc["_confianca"] = "alta"
        else:
            doc["_confianca"] = "media"
            doc["_tipo_alternativo"] = tipo_correto
            doc["_razao_dupla_checagem"] = (dados.get("razao") or "")[:200]
    except Exception:
        pass  # se falhar, nao bloqueia o fluxo

# Sequencia para INSS Administrativo
# REGRA (definida pelo time Jaine/Andre):
# - grupo_procuracao_termos = procuracao + substabelecimento + termo_representacao + protocolo
#   (APENAS esses 4 tipos no mesmo arquivo)
# - DECLARACOES (residencia, hipossuficiencia, beneficios INSS) ficam SEPARADAS
# - Documentos do INSS (comunicacao decisao, declaracao beneficios) ficam SEPARADOS
# - Contrato de Honorarios SEMPRE separado (sigiloso, nao vai pro processo)
SEQUENCIA_INSS_ADMIN = [
    ("grupo_procuracao_termos", "Procuracao_Termos_Assinaturas"),
    ("declaracao_hipossuficiencia", "Declaracao_Hipossuficiencia"),
    ("declaracao", "Declaracao"),
    ("contrato_de_honorarios", "Contrato_de_Honorarios_SIGILOSO"),
    ("documento_pessoal", None),  # usa tipo real
    ("carta_indeferimento", "Comunicacao_Decisao_INSS"),
    ("declaracao_beneficios_inss", "Declaracao_Beneficios_INSS"),
]

# Sequencia detalhada para Processos Judiciais
# Ordem definida pelo time juridico (Andre, 29/04/2026):
SEQUENCIA_JUDICIAL = [
    ("peticao_inicial", "Peticao_Inicial"),                                      # 01
    ("calculo_valor_causa", "Calculo_Valor_da_Causa"),                           # 02
    ("procuracao", "Procuracao"),                                                # 03
    ("substabelecimento", "Substabelecimento"),                                  # 04
    ("declaracao_hipossuficiencia", "Declaracao_de_Hipossuficiencia"),           # 05
    ("documento_pessoal", "RG_do_Autor"),                                        # 06
    ("carta_indeferimento", "Carta_Indeferimento_INSS"),                         # 07
    ("cnis", "CNIS"),                                                            # 08
    ("ctps", "CTPS"),                                                            # 09
    ("atestados_relatorios_receitas", "Atestados_Relatorios_Receitas_Medicas"),  # 10
    ("exames_medicos", "Exames_Medicos"),                                        # 11
    ("laudo_meuinss", "Laudo_MeuINSS"),                                          # 12
    ("documentos_rurais", "Documentos_Rurais"),                                  # 13
    ("folha_v7", "Folha_V7"),                                                    # 14
    ("certidoes", "Certidoes"),                                                  # 15
    ("termo_homologacao_rural", "Termo_Homologacao_Atividade_Rural"),            # 16 (inclui Despacho Decisorio)
    ("declaracao_tempo_servico", "Declaracao_Tempo_Servico_Ficha_Funcionario"),  # 17
    ("certidao_tempo_servico", "Certidao_Tempo_Servico"),                        # 18
    ("ficha_financeira", "Ficha_Financeira"),                                    # 19
    ("ppp", "PPP"),                                                              # 20
    ("ltcat", "LTCAT"),                                                          # 21
    ("gps", "GPS"),                                                              # 22
    ("comprovante_gasto", "Comprovantes_Gastos"),                                # 23
    ("foto_residencia", "Fotos_Residencia"),                                     # 24
    ("avaliacao_social_pericia", "Avaliacao_Social_Pericia_Medica"),             # 25 (Quadro de Informacoes)
    ("contagem_tempo", "Contagem_Tempo"),                                        # 26
    ("calculo_regras_transicao", "Calculo_Regras_Transicao"),                    # 27
    ("calculo_rmi", "Calculo_RMI"),                                              # 28
    ("copia_processo_administrativo", "Copia_Processo_Administrativo"),          # 29
    ("certidao_negativa_estadual", "Certidao_Negativa_Justica_Estadual"),        # 30
    ("comprovante_residencia", "Comprovante_Residencia"),                        # 31
]

# Categorias que devem ser MERGED em um unico PDF (em ordem cronologica)
CATEGORIAS_MERGE = {
    "atestados_relatorios_receitas",
    "exames_medicos",
    "documentos_rurais",
    "comprovante_gasto",
    "ctps",
    "certidoes",
    "declaracao_tempo_servico",
    "ficha_financeira",
    "gps",
    "grupo_procuracao_termos",  # admin: procuracao+substabelecimento+termo juntos
}

# Categorias que devem receber o "protocolo de assinatura" anexado
CATEGORIAS_ASSINADAS = {
    "procuracao",
    "substabelecimento",
    "declaracao",
    "declaracao_hipossuficiencia",
    "contrato_de_honorarios",
    "termo_de_responsabilidade",
}

# Para INSS admin: agrupar EXCLUSIVAMENTE estes tipos em UM UNICO arquivo
# (procuracoes + substabelecimentos + termos + protocolos de assinatura)
# NAO entra aqui:
#   - declaracoes de qualquer tipo (residencia, hipossuficiencia, beneficios)
#   - contrato_de_honorarios (sigiloso)
#   - qualquer documento do INSS (comunicacao decisao, declaracao beneficios)
GRUPO_MERGE_ADMIN = {
    "grupo_procuracao_termos": [
        "procuracao",
        "substabelecimento",
        "termo_de_responsabilidade",
        "protocolo_assinatura",
    ],
}


def limpar_nome(texto):
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = texto.lower().replace(" ", "_")
    texto = "".join(c for c in texto if c.isalnum() or c in ("_", "-"))
    return texto


def contar_paginas_pdf(caminho):
    """Conta total de paginas de um PDF sem extrair texto."""
    try:
        from pypdf import PdfReader
        return len(PdfReader(caminho).pages)
    except Exception:
        return 0


def extrair_textos_todas_paginas(caminho):
    """Extrai texto de todas as paginas de um PDF.
    Retorna dict com paginas + aviso caso tenha truncado."""
    import pdfplumber
    resultado = {"paginas": [], "total_real": 0, "truncado": False}
    try:
        with pdfplumber.open(caminho) as pdf:
            total_real = len(pdf.pages)
            resultado["total_real"] = total_real
            total = min(total_real, MAX_PAGINAS_PDF)
            if total_real > MAX_PAGINAS_PDF:
                resultado["truncado"] = True
            # PDFs grandes: menos texto por pagina para economizar memoria e tokens
            if total > CHUNK_SIZE_PAGINAS:
                max_linhas = 15
                max_chars = 1000
            else:
                max_linhas = 20
                max_chars = MAX_TEXTO_POR_PAGINA
            for i in range(total):
                try:
                    texto = pdf.pages[i].extract_text() or ""
                except Exception:
                    texto = ""
                linhas = texto.split("\n")[:max_linhas]
                resultado["paginas"].append({"pagina": i + 1, "texto": "\n".join(linhas)[:max_chars]})
        return resultado
    except Exception:
        return resultado


def dividir_pdf_por_tamanho(caminho, tmp_dir, max_bytes=MAX_FILE_SIZE_BYTES):
    """Divide um PDF grande em partes menores de no maximo max_bytes cada."""
    from pypdf import PdfReader, PdfWriter

    tamanho = os.path.getsize(caminho)
    if tamanho <= max_bytes:
        return [caminho]  # nao precisa dividir

    try:
        reader = PdfReader(caminho)
        total_paginas = len(reader.pages)
        if total_paginas <= 1:
            return [caminho]  # nao da pra dividir mais

        # Estima paginas por parte baseado no tamanho
        bytes_por_pagina = tamanho / total_paginas
        paginas_por_parte = max(1, int(max_bytes / bytes_por_pagina))

        partes = []
        inicio = 0
        parte_num = 1

        while inicio < total_paginas:
            fim = min(inicio + paginas_por_parte, total_paginas)
            writer = PdfWriter()
            for i in range(inicio, fim):
                writer.add_page(reader.pages[i])

            nome_base = Path(caminho).stem
            parte_path = os.path.join(tmp_dir, f"{nome_base}_parte{parte_num}.pdf")
            with open(parte_path, "wb") as f:
                writer.write(f)

            # Verifica se a parte ainda e grande demais (pode acontecer com paginas pesadas)
            if os.path.getsize(parte_path) > max_bytes and (fim - inicio) > 1:
                # Tenta com menos paginas
                os.remove(parte_path)
                paginas_por_parte = max(1, paginas_por_parte // 2)
                continue

            partes.append(parte_path)
            inicio = fim
            parte_num += 1

        return partes if partes else [caminho]
    except Exception:
        return [caminho]


def pagina_para_imagem_b64(caminho, pagina_num):
    """Converte uma pagina especifica do PDF em imagem base64."""
    from pdf2image import convert_from_path
    try:
        imagens = convert_from_path(caminho, first_page=pagina_num, last_page=pagina_num, dpi=150)
        if imagens:
            buffer = io.BytesIO()
            imagens[0].save(buffer, format="JPEG", quality=80)
            return base64.b64encode(buffer.getvalue()).decode()
        return None
    except Exception:
        return None


def separar_pdf(caminho_original, pagina_inicio, pagina_fim, caminho_destino):
    """Extrai paginas de um PDF e salva em novo arquivo."""
    import pdfplumber
    from pypdf import PdfReader, PdfWriter
    try:
        reader = PdfReader(caminho_original)
        writer = PdfWriter()
        for i in range(pagina_inicio - 1, min(pagina_fim, len(reader.pages))):
            writer.add_page(reader.pages[i])
        with open(caminho_destino, "wb") as f:
            writer.write(f)
        return True
    except Exception:
        return False


def _parse_json_response(resposta):
    """Parse JSON da resposta da IA, lidando com markdown code blocks."""
    resposta = resposta.strip()
    if resposta.startswith("{"):
        return json.loads(resposta)
    if "```" in resposta:
        json_str = resposta.split("```")[1]
        if json_str.startswith("json"):
            json_str = json_str[4:]
        return json.loads(json_str.strip())
    return json.loads(resposta)


def _mapear_chunk(client, caminho, nome_arquivo, chunk_paginas):
    """Mapeia documentos em um chunk de paginas (max CHUNK_SIZE_PAGINAS)."""
    texto_chunk = ""
    for p in chunk_paginas:
        texto_chunk += f"\n--- PAGINA {p['pagina']} ---\n{p['texto']}\n"

    # Para paginas sem texto no chunk, tenta imagem (max 1 por chunk para economizar)
    imagens_content = []
    paginas_sem_texto = [p for p in chunk_paginas if not p["texto"].strip()]
    if paginas_sem_texto and len(paginas_sem_texto) <= 2:
        p = paginas_sem_texto[0]
        img_b64 = pagina_para_imagem_b64(caminho, p["pagina"])
        if img_b64:
            imagens_content.append({"type": "text", "text": f"--- PAGINA {p['pagina']} (imagem) ---"})
            imagens_content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64},
            })
            img_b64 = None

    pag_inicio = chunk_paginas[0]["pagina"]
    pag_fim = chunk_paginas[-1]["pagina"]

    messages_content = [{"type": "text",
        "text": f"Arquivo: {nome_arquivo} (paginas {pag_inicio}-{pag_fim})\n\n{texto_chunk}"}]
    messages_content.extend(imagens_content)

    try:
        response = client.messages.create(
            model=MODELO_IA,
            max_tokens=800,
            messages=[{"role": "user", "content": messages_content + [
                {"type": "text", "text": PROMPT_MAPEAMENTO_CURTO}
            ]}],
        )
        dados = _parse_json_response(response.content[0].text)
        docs = dados.get("documentos", [])
        if docs:
            return docs
    except Exception:
        pass

    # Fallback: chunk inteiro como documento unico
    return [{"tipo": "desconhecido", "pagina_inicio": pag_inicio, "pagina_fim": pag_fim, "data": None}]


def _merge_boundary_docs(docs):
    """Junta documentos do mesmo tipo que ficaram separados na fronteira entre chunks."""
    if len(docs) <= 1:
        return docs
    merged = [docs[0]]
    for doc in docs[1:]:
        prev = merged[-1]
        # Mesmo tipo e paginas consecutivas? Provavelmente o mesmo documento
        if (limpar_nome(doc.get("tipo", "")) == limpar_nome(prev.get("tipo", "")) and
                doc.get("pagina_inicio", 0) == prev.get("pagina_fim", 0) + 1):
            prev["pagina_fim"] = doc["pagina_fim"]
            if doc.get("data") and not prev.get("data"):
                prev["data"] = doc["data"]
        else:
            merged.append(doc)
    return merged


def mapear_documentos_pdf(client, caminho, nome_arquivo):
    """Analisa PDF com multiplas paginas e identifica cada documento separado."""
    dados_extracao = extrair_textos_todas_paginas(caminho)
    paginas = dados_extracao["paginas"]
    truncado = dados_extracao.get("truncado", False)
    total_real = dados_extracao.get("total_real", 0)

    if not paginas:
        # PDF escaneado - tenta via imagem da primeira pagina
        img_b64 = pagina_para_imagem_b64(caminho, 1)
        if img_b64:
            resultado = analisar_imagem_simples(client, img_b64, nome_arquivo)
            return [{"tipo": resultado.get("tipo", "desconhecido"), "pagina_inicio": 1,
                     "pagina_fim": 1, "data": resultado.get("data")}]
        return [{"tipo": "desconhecido", "pagina_inicio": 1, "pagina_fim": 1, "data": None}]

    total_paginas = len(paginas)

    # Se PDF foi truncado, avisa nos resultados (pagina extra invisivel que o usuario pode ver no relatorio)
    aviso_truncado = None
    if truncado:
        aviso_truncado = f"ATENCAO: PDF tem {total_real} paginas, so foram analisadas as primeiras {MAX_PAGINAS_PDF}. Divida o arquivo em partes menores."

    if total_paginas == 1:
        # PDF de 1 pagina - analise simples
        texto = paginas[0]["texto"]
        if texto.strip():
            resultado = analisar_texto_simples(client, texto, nome_arquivo)
        else:
            img_b64 = pagina_para_imagem_b64(caminho, 1)
            resultado = analisar_imagem_simples(client, img_b64, nome_arquivo) if img_b64 else {"tipo": "desconhecido", "data": None}
        doc = {"tipo": resultado.get("tipo", "desconhecido"), "pagina_inicio": 1,
                 "pagina_fim": 1, "data": resultado.get("data")}
        if aviso_truncado:
            doc["_aviso"] = aviso_truncado
        return [doc]

    # === PDF escaneado (>= 70% das paginas sem texto): processa pagina por pagina via Vision ===
    # Caso do "PRINT ...jpg" ou DOCS MEDICOS com fotos — pdfplumber nao extrai texto.
    paginas_sem_texto = sum(1 for p in paginas if not p["texto"].strip())
    if paginas_sem_texto >= total_paginas * 0.7 and total_paginas <= 15:
        docs_scan = []
        for p in paginas:
            img_b64 = pagina_para_imagem_b64(caminho, p["pagina"])
            if not img_b64:
                continue
            resultado = analisar_imagem_simples(client, img_b64, nome_arquivo)
            img_b64 = None  # libera memoria
            docs_scan.append({
                "tipo": resultado.get("tipo", "desconhecido"),
                "pagina_inicio": p["pagina"],
                "pagina_fim": p["pagina"],
                "data": resultado.get("data"),
            })
        if docs_scan:
            # Junta paginas consecutivas do mesmo tipo (ex: laudo medico de 3 paginas)
            docs_scan = _merge_boundary_docs(docs_scan)
            if aviso_truncado:
                docs_scan[0]["_aviso"] = aviso_truncado
            return docs_scan

    # === PDF pequeno (ate CHUNK_SIZE paginas): chamada unica (comportamento original) ===
    if total_paginas <= CHUNK_SIZE_PAGINAS:
        texto_completo = ""
        for p in paginas:
            texto_completo += f"\n--- PAGINA {p['pagina']} ---\n{p['texto']}\n"

        paginas_sem_texto = [p for p in paginas if not p["texto"].strip()]
        imagens_content = []
        if paginas_sem_texto and len(paginas_sem_texto) <= 3:
            for p in paginas_sem_texto[:2]:
                img_b64 = pagina_para_imagem_b64(caminho, p["pagina"])
                if img_b64:
                    imagens_content.append({"type": "text", "text": f"--- PAGINA {p['pagina']} (imagem) ---"})
                    imagens_content.append({
                        "type": "image",
                        "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64},
                    })
                    img_b64 = None

        messages_content = [{"type": "text",
            "text": f"Arquivo: {nome_arquivo} ({total_paginas} paginas)\n\n{texto_completo}"}]
        messages_content.extend(imagens_content)

        try:
            response = client.messages.create(
                model=MODELO_IA,
                max_tokens=1000,
                messages=[
                    {"role": "user", "content": messages_content},
                    {"role": "user", "content": PROMPT_MAPEAMENTO},
                ],
            )
            dados = _parse_json_response(response.content[0].text)
            docs = dados.get("documentos", [])
            if docs:
                if aviso_truncado:
                    docs[0]["_aviso"] = aviso_truncado
                return docs
        except Exception:
            pass

        doc = {"tipo": "desconhecido", "pagina_inicio": 1, "pagina_fim": total_paginas, "data": None}
        if aviso_truncado:
            doc["_aviso"] = aviso_truncado
        return [doc]

    # === PDF grande: processa em chunks de CHUNK_SIZE_PAGINAS paginas ===
    todos_docs = []
    for i in range(0, total_paginas, CHUNK_SIZE_PAGINAS):
        chunk = paginas[i:i + CHUNK_SIZE_PAGINAS]
        chunk_docs = _mapear_chunk(client, caminho, nome_arquivo, chunk)
        todos_docs.extend(chunk_docs)

    # Junta documentos que ficaram divididos na fronteira entre chunks
    todos_docs = _merge_boundary_docs(todos_docs)

    if not todos_docs:
        todos_docs = [{"tipo": "desconhecido", "pagina_inicio": 1, "pagina_fim": total_paginas, "data": None}]

    if aviso_truncado:
        todos_docs[0]["_aviso"] = aviso_truncado

    return todos_docs


def analisar_texto_simples(client, texto, nome_arquivo):
    """Analise simples de texto para extrair tipo e data."""
    try:
        response = client.messages.create(
            model=MODELO_IA,
            max_tokens=200,
            messages=[
                {"role": "user", "content": f"Nome do arquivo: {nome_arquivo}\n\nConteudo:\n{texto}"},
                {"role": "user", "content": PROMPT_EXTRACAO_SIMPLES},
            ],
        )
        resposta = response.content[0].text.strip()
        if resposta.startswith("{"):
            return json.loads(resposta)
        if "```" in resposta:
            json_str = resposta.split("```")[1]
            if json_str.startswith("json"):
                json_str = json_str[4:]
            return json.loads(json_str.strip())
        return json.loads(resposta)
    except Exception:
        return {"tipo": "desconhecido", "data": None}


def analisar_imagem_simples(client, imagem_b64, nome_arquivo):
    """Analise simples de imagem para extrair tipo e data."""
    try:
        response = client.messages.create(
            model=MODELO_IA,
            max_tokens=200,
            messages=[
                {"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": imagem_b64}},
                    {"type": "text", "text": f"Nome do arquivo: {nome_arquivo}"},
                ]},
                {"role": "user", "content": PROMPT_EXTRACAO_SIMPLES},
            ],
        )
        resposta = response.content[0].text.strip()
        if resposta.startswith("{"):
            return json.loads(resposta)
        return json.loads(resposta)
    except Exception:
        return {"tipo": "desconhecido", "data": None}


# ============ Extracao de data com fallbacks em cascata (#3.3) ============

# Regex para datas comuns em nomes de arquivo
# Aceita: 2023-03-15, 2023_03_15, 15-03-2023, 15/03/2023, 15.03.2023, 20230315
_RE_DATAS_FILENAME = [
    # YYYY-MM-DD ou YYYY_MM_DD ou YYYY.MM.DD
    (re.compile(r"(20\d{2})[-_./](\d{1,2})[-_./](\d{1,2})"), "ymd"),
    # DD-MM-YYYY ou DD/MM/YYYY
    (re.compile(r"(\d{1,2})[-_./](\d{1,2})[-_./](20\d{2})"), "dmy"),
    # YYYYMMDD compactado
    (re.compile(r"(20\d{2})(\d{2})(\d{2})"), "ymd"),
]


def extrair_data_pdf_metadata(caminho):
    """Tenta ler a data de criacao do PDF dos metadados. Retorna YYYY-MM-DD ou None."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(caminho)
        meta = reader.metadata
        if not meta:
            return None
        # Pega a data mais antiga entre CreationDate e ModDate (geralmente mais confiavel)
        candidatos = []
        for campo in ["/CreationDate", "/ModDate"]:
            valor = meta.get(campo)
            if valor:
                candidatos.append(str(valor))
        for raw in candidatos:
            # Formato pdf: "D:20230315143022-03'00'" ou "D:20230315143022Z"
            m = re.search(r"(20\d{2})(\d{2})(\d{2})", raw)
            if m:
                ano, mes, dia = m.groups()
                # Valida mes/dia basicamente
                if 1 <= int(mes) <= 12 and 1 <= int(dia) <= 31:
                    return f"{ano}-{mes}-{dia}"
        return None
    except Exception:
        return None


def extrair_data_filename(nome_arquivo):
    """Tenta extrair uma data do nome do arquivo. Retorna YYYY-MM-DD ou None."""
    if not nome_arquivo:
        return None
    nome = Path(nome_arquivo).stem  # remove extensao
    for regex, formato in _RE_DATAS_FILENAME:
        m = regex.search(nome)
        if not m:
            continue
        g1, g2, g3 = m.groups()
        if formato == "ymd":
            ano, mes, dia = g1, g2.zfill(2), g3.zfill(2)
        else:  # dmy
            dia, mes, ano = g1.zfill(2), g2.zfill(2), g3
        # Validacao basica
        try:
            if 1 <= int(mes) <= 12 and 1 <= int(dia) <= 31 and 2000 <= int(ano) <= 2100:
                return f"{ano}-{mes}-{dia}"
        except ValueError:
            continue
    return None


def resolver_data_fallback(resultado, caminho_original, nome_original):
    """Aplica fallbacks em cascata pra tentar encontrar uma data quando a IA nao achou.
    Modifica o dicionario 'resultado' in-place adicionando:
    - data: YYYY-MM-DD ou None
    - _data_fonte: "ia" | "pdf_metadata" | "filename" | None
    """
    # Se IA ja achou a data, marca como tal e retorna
    if resultado.get("data"):
        resultado["_data_fonte"] = "ia"
        return

    # Fallback 1: metadata do PDF
    if caminho_original and Path(caminho_original).suffix.lower() == ".pdf":
        data_meta = extrair_data_pdf_metadata(caminho_original)
        if data_meta:
            resultado["data"] = data_meta
            resultado["_data_fonte"] = "pdf_metadata"
            return

    # Fallback 2: nome do arquivo original
    data_nome = extrair_data_filename(nome_original)
    if data_nome:
        resultado["data"] = data_nome
        resultado["_data_fonte"] = "filename"
        return

    # Sem data
    resultado["_data_fonte"] = None


def processar_arquivo_completo(client, caminho, nome_original, tmp_dir):
    """Processa um arquivo, separando-o em documentos individuais se necessario.
    Retorna lista de resultados (1 por documento encontrado)."""
    ext = Path(nome_original).suffix.lower()
    resultados = []

    if ext == ".pdf":
        # Mapeia todos os documentos dentro do PDF
        docs_encontrados = mapear_documentos_pdf(client, caminho, nome_original)

        import pdfplumber
        try:
            with pdfplumber.open(caminho) as pdf:
                total_paginas = len(pdf.pages)
        except Exception:
            total_paginas = 1

        if len(docs_encontrados) == 1 and total_paginas <= 2:
            # Documento unico, usa o arquivo original
            doc = docs_encontrados[0]
            resultados.append({
                "tipo": doc.get("tipo", "desconhecido"),
                "data": doc.get("data"),
                "arquivo_tmp": caminho,
                "nome_original": nome_original,
                "extensao": ".pdf",
                "_aviso": doc.get("_aviso"),
            })
        else:
            # Multiplos documentos - separa em arquivos individuais
            for i, doc in enumerate(docs_encontrados):
                p_inicio = doc.get("pagina_inicio", 1)
                p_fim = doc.get("pagina_fim", p_inicio)
                tipo = doc.get("tipo", "desconhecido")
                tipo_limpo = limpar_nome(tipo)[:30]

                novo_caminho = os.path.join(tmp_dir, f"split_{i}_{tipo_limpo}.pdf")
                if separar_pdf(caminho, p_inicio, p_fim, novo_caminho):
                    resultados.append({
                        "tipo": tipo,
                        "data": doc.get("data"),
                        "arquivo_tmp": novo_caminho,
                        "nome_original": f"{nome_original} (pag {p_inicio}-{p_fim})",
                        "extensao": ".pdf",
                        "_aviso": doc.get("_aviso") if i == 0 else None,
                    })

    elif ext in (".jpg", ".jpeg", ".png"):
        with open(caminho, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode()
        resultado = analisar_imagem_simples(client, img_b64, nome_original)
        resultados.append({
            "tipo": resultado.get("tipo", "desconhecido"),
            "data": resultado.get("data"),
            "arquivo_tmp": caminho,
            "nome_original": nome_original,
            "extensao": ext,
        })

    elif ext == ".docx":
        from docx import Document
        try:
            doc = Document(caminho)
            texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if texto:
                resultado = analisar_texto_simples(client, texto[:4000], nome_original)
            else:
                resultado = {"tipo": "desconhecido", "data": None}
        except Exception:
            resultado = {"tipo": "desconhecido", "data": None}

        resultados.append({
            "tipo": resultado.get("tipo", "desconhecido"),
            "data": resultado.get("data"),
            "arquivo_tmp": caminho,
            "nome_original": nome_original,
            "extensao": ext,
        })

    if not resultados:
        resultados.append({
            "tipo": "desconhecido",
            "data": None,
            "arquivo_tmp": caminho,
            "nome_original": nome_original,
            "extensao": ext,
        })

    return resultados


def eh_protocolo_assinatura(r):
    """Verifica se um documento e protocolo de assinatura."""
    tipo = limpar_nome(r.get("tipo", ""))
    return "protocolo" in tipo and "assinatura" in tipo


def classificar_doc(r, tipo_processo):
    """Retorna a categoria do doc para ordenacao baseada no tipo de processo."""
    tipo = limpar_nome(r.get("tipo", ""))

    # === DOCS JUDICIAIS NOVOS (precedencia maxima) ===
    # Peticao Inicial
    if "peticao" in tipo and "inicial" in tipo:
        return "peticao_inicial"
    if tipo == "inicial":
        return "peticao_inicial"
    # Calculo do valor da causa
    if "calculo" in tipo and "valor" in tipo and "causa" in tipo:
        return "calculo_valor_causa"
    if "valor" in tipo and "causa" in tipo:
        return "calculo_valor_causa"
    # Despacho decisorio (segundo Andre, vai com Termo de Homologacao Rural)
    if "despacho" in tipo and ("decisorio" in tipo or "decisao" in tipo):
        return "termo_homologacao_rural"
    # Quadro de informacoes da avaliacao social / pericia medica
    if "quadro" in tipo and ("avaliacao" in tipo or "pericia" in tipo or "informacoes" in tipo):
        return "avaliacao_social_pericia"

    # Protocolo de assinatura digital (inclui "Relatorio de Assinaturas" do ZapSign, etc)
    if "protocolo" in tipo and "assinatura" in tipo:
        return "protocolo_assinatura"
    if "relatorio" in tipo and "assinatura" in tipo:
        return "protocolo_assinatura"
    # Substabelecimento (separado de procuracao)
    if "substabelecimento" in tipo:
        return "substabelecimento"
    # Procuracao
    if "procuracao" in tipo:
        return "procuracao"
    # === DOCUMENTOS DO INSS (NUNCA entram no grupo admin, ficam separados) ===
    # Declaracao de Beneficios INSS (documento emitido pelo INSS)
    if "declaracao" in tipo and ("beneficio" in tipo or "inss" in tipo):
        return "declaracao_beneficios_inss"
    if "declaracao" in tipo and "sistema_unico" in tipo:
        return "declaracao_beneficios_inss"
    # Comunicacao de Decisao INSS
    if "comunicacao" in tipo and "decisao" in tipo:
        return "carta_indeferimento"
    # Carta de indeferimento INSS
    if "carta" in tipo and "indeferimento" in tipo:
        return "carta_indeferimento"
    if "decisao" in tipo and "inss" in tipo:
        return "carta_indeferimento"
    if ("indeferimento" in tipo or "negado" in tipo) and ("inss" in tipo or "beneficio" in tipo):
        return "carta_indeferimento"
    # === DECLARACOES DO CLIENTE (ficam separadas, nao vao pro grupo admin) ===
    # Declaracao de hipossuficiencia (tem categoria propria)
    if "declaracao" in tipo and ("hipossuficiencia" in tipo or "pobreza" in tipo):
        return "declaracao_hipossuficiencia"
    # Outras declaracoes
    if "declaracao" in tipo and "tempo" in tipo and "servico" in tipo:
        return "declaracao_tempo_servico"
    if "declaracao" in tipo and "residencia" in tipo:
        return "declaracao"  # declaracao de residencia do cliente
    if "declaracao" in tipo:
        return "declaracao"
    # Contrato de honorarios
    if "contrato" in tipo and "honorario" in tipo:
        return "contrato_de_honorarios"
    # Termo de responsabilidade / representacao
    if "termo" in tipo and ("responsabilidade" in tipo or "representacao" in tipo):
        return "termo_de_responsabilidade"
    # CNIS
    if "cnis" in tipo:
        return "cnis"
    # CTPS
    if "ctps" in tipo or ("carteira" in tipo and "trabalho" in tipo):
        return "ctps"
    # Documentos medicos
    if "laudo" in tipo and ("meuinss" in tipo or "meu_inss" in tipo):
        return "laudo_meuinss"
    if "exame" in tipo and "medico" in tipo:
        return "exames_medicos"
    if "exame" in tipo:
        return "exames_medicos"
    if any(k in tipo for k in ["atestado", "relatorio_medico", "receita_medica"]):
        return "atestados_relatorios_receitas"
    if "laudo" in tipo:
        return "atestados_relatorios_receitas"
    # Documentos rurais
    if "rural" in tipo and "homologacao" in tipo:
        return "termo_homologacao_rural"
    if "rural" in tipo or "atividade_rural" in tipo:
        return "documentos_rurais"
    # Folha V7
    if "folha_v7" in tipo or tipo == "folha_v7" or "v7" in tipo:
        return "folha_v7"
    # Certidoes
    if "certidao" in tipo and "tempo" in tipo and "servico" in tipo:
        return "certidao_tempo_servico"
    if "certidao" in tipo and "negativa" in tipo:
        return "certidao_negativa_estadual"
    if "certidao" in tipo and any(k in tipo for k in ["casamento", "nascimento", "obito"]):
        return "certidoes"
    if "certidao" in tipo:
        return "certidoes"
    # Ficha financeira / funcionario
    if "ficha" in tipo and "financeira" in tipo:
        return "ficha_financeira"
    if "ficha" in tipo and "funcionario" in tipo:
        return "declaracao_tempo_servico"
    # PPP / LTCAT
    if "ppp" in tipo or "perfil_profissiografico" in tipo:
        return "ppp"
    if "ltcat" in tipo or ("laudo" in tipo and "tecnico" in tipo):
        return "ltcat"
    # GPS
    if "gps" in tipo or ("guia" in tipo and "previdencia" in tipo):
        return "gps"
    # Comprovantes
    if "comprovante" in tipo and ("gasto" in tipo or "despesa" in tipo or "pagamento" in tipo):
        return "comprovante_gasto"
    if "comprovante" in tipo and "residencia" in tipo:
        return "comprovante_residencia"
    # Fotos
    if "foto" in tipo and "residencia" in tipo:
        return "foto_residencia"
    # Avaliacao social / pericia
    if ("avaliacao" in tipo and "social" in tipo) or "pericia_medica" in tipo or "pericia" in tipo:
        return "avaliacao_social_pericia"
    # Calculos
    if "contagem" in tipo and "tempo" in tipo:
        return "contagem_tempo"
    if "calculo" in tipo and ("transicao" in tipo or "regra" in tipo):
        return "calculo_regras_transicao"
    if "calculo" in tipo and ("rmi" in tipo or "renda_mensal" in tipo):
        return "calculo_rmi"
    # Copia processo
    if "copia" in tipo and "processo" in tipo:
        return "copia_processo_administrativo"
    if "processo_administrativo" in tipo:
        return "copia_processo_administrativo"
    # Documentos pessoais
    if tipo in ("rg", "cpf", "cnh") or \
       any(k in tipo for k in ["identidade", "cnh", "carteira_nacional"]):
        return "documento_pessoal"

    return None


def merge_pdfs(caminhos, destino):
    """Junta varios PDFs em um unico arquivo."""
    from pypdf import PdfReader, PdfWriter
    try:
        writer = PdfWriter()
        for caminho in caminhos:
            try:
                reader = PdfReader(caminho)
                for page in reader.pages:
                    writer.add_page(page)
            except Exception:
                continue
        with open(destino, "wb") as f:
            writer.write(f)
        return True
    except Exception:
        return False


# === ROTAS ===

@app.route("/")
def index():
    return render_template("index.html", tipos=TIPOS_PROCESSO, usuarios=USUARIOS)


def registrar_uso(usuario, nome_cliente, tipo_processo, total_docs, status):
    """Envia log de uso/erro para webhook (n8n -> Google Sheets).
    Erros sao diferenciados pelo prefixo 'ERRO[tipo_erro]' no campo status.
    Tambem loga em stdout (Render captura por 7 dias) pra ter persistencia mesmo se o n8n falhar.
    """
    # 1) LOG LOCAL (stdout — Render captura) — funciona sempre, ate sem webhook
    audit_log.info(
        "USO | usuario=%s | cliente=%s | tipo=%s | docs=%s | status=%s",
        usuario or "desconhecido",
        nome_cliente or "",
        TIPOS_PROCESSO.get(tipo_processo, tipo_processo or ""),
        total_docs,
        status,
    )

    # 1.5) PERSISTENCIA LOCAL (SQLite) — sobrevive restart do worker
    evento = "ERRO" if (status or "").startswith("ERRO[") or (status or "").startswith("erro:") else "SUCESSO"
    audit_save(
        evento=evento,
        usuario=usuario,
        cliente=nome_cliente,
        tipo_processo=TIPOS_PROCESSO.get(tipo_processo, tipo_processo or ""),
        total_documentos=total_docs,
        mensagem=status,
    )

    # 2) WEBHOOK (n8n -> Google Sheets) — pode falhar silenciosamente
    if not LOG_WEBHOOK_URL:
        return
    try:
        import urllib.request
        payload = json.dumps({
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "usuario": usuario or "desconhecido",
            "cliente": nome_cliente or "",
            "tipo_processo": TIPOS_PROCESSO.get(tipo_processo, tipo_processo or ""),
            "total_documentos": str(total_docs),
            "status": status,
        }).encode("utf-8")
        req = urllib.request.Request(
            LOG_WEBHOOK_URL,
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        urllib.request.urlopen(req, timeout=5)
    except Exception:
        pass  # nao quebra o app se o log falhar


def registrar_erro(usuario, nome_cliente, tipo_processo, arquivo, tipo_erro, mensagem):
    """Registra um erro na MESMA planilha de uso, com prefixo ERRO[tipo_erro].
    Tambem loga em stdout pra auditoria local (Render captura por 7 dias).
    """
    msg_limpa = str(mensagem)[:300].replace("\n", " ").replace("|", "/")
    arq = arquivo or "-"

    # LOG LOCAL detalhado (separado do registrar_uso pra ter mais info no log)
    audit_log.error(
        "ERRO[%s] | usuario=%s | cliente=%s | tipo=%s | arquivo=%s | msg=%s",
        tipo_erro,
        usuario or "desconhecido",
        nome_cliente or "",
        TIPOS_PROCESSO.get(tipo_processo, tipo_processo or ""),
        arq,
        msg_limpa,
    )

    # PERSISTENCIA LOCAL — informacao detalhada do erro vai pro SQLite
    audit_save(
        evento="ERRO",
        usuario=usuario,
        cliente=nome_cliente,
        tipo_processo=TIPOS_PROCESSO.get(tipo_processo, tipo_processo or ""),
        arquivos=arquivo,
        tipo_erro=tipo_erro,
        mensagem=msg_limpa,
    )

    # Tambem escreve na planilha (via webhook) — pode falhar silenciosamente
    status_codificado = f"ERRO[{tipo_erro}] {arq}: {msg_limpa}"
    registrar_uso(usuario, nome_cliente, tipo_processo, 0, status_codificado)


@app.route("/processar", methods=["POST"])
def processar():
    import anthropic

    # Limpa ZIPs antigos (>1h) a cada request para nao entupir o disco
    limpar_zips_antigos()

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return jsonify({"erro": "API key nao configurada no servidor"}), 500

    usuario = request.form.get("usuario", "").strip()
    nome_cliente = request.form.get("nome_cliente", "").strip()
    tipo_processo = request.form.get("tipo_processo", "").strip()

    if not usuario or usuario not in USUARIOS:
        return jsonify({"erro": "Selecione o usuario que esta processando"}), 400
    if not nome_cliente:
        return jsonify({"erro": "Nome do cliente e obrigatorio"}), 400
    if tipo_processo not in TIPOS_PROCESSO:
        return jsonify({"erro": "Tipo de processo invalido"}), 400

    arquivos = request.files.getlist("documentos")
    if not arquivos or all(f.filename == "" for f in arquivos):
        return jsonify({"erro": "Nenhum documento enviado"}), 400

    arquivos_validos = [
        f for f in arquivos
        if f.filename and Path(f.filename).suffix.lower() in EXTENSOES_ACEITAS
    ]
    if not arquivos_validos:
        return jsonify({"erro": f"Nenhum arquivo valido. Aceitos: {', '.join(EXTENSOES_ACEITAS)}"}), 400

    tmp_dir = tempfile.mkdtemp()

    try:
        client = anthropic.Anthropic(api_key=api_key)
        resultados = []

        for arquivo in arquivos_validos:
            nome_original = arquivo.filename
            ext = Path(nome_original).suffix.lower()
            tmp_path = os.path.join(tmp_dir, nome_original)
            arquivo.save(tmp_path)

            try:
                # Processa e separa documentos automaticamente
                docs_separados = processar_arquivo_completo(client, tmp_path, nome_original, tmp_dir)
                # Aplica fallback de data em cada documento (metadata PDF / nome do arquivo)
                for doc in docs_separados:
                    resolver_data_fallback(doc, tmp_path, nome_original)
                resultados.extend(docs_separados)
            except Exception as e:
                # Nao falha o request inteiro — adiciona como documento desconhecido
                resultados.append({
                    "tipo": "desconhecido",
                    "data": None,
                    "arquivo_tmp": tmp_path,
                    "nome_original": nome_original,
                    "extensao": ext,
                    "_data_fonte": None,
                })
                registrar_erro(usuario, nome_cliente, tipo_processo, nome_original,
                              "excecao_processamento_arquivo", e)

        # ===== ETAPA 1: Anexar protocolos de assinatura ao documento anterior =====
        # Para JUDICIAL/consumidor/trabalhista/civel: protocolo anexa ao doc assinado anterior
        # Para INSS ADMIN: protocolo vai direto pro grupo_procuracao_termos (nao anexa)
        resultados_processados = []
        protocolos_pendentes = []

        for r in resultados:
            if eh_protocolo_assinatura(r) and tipo_processo != "inss_admin":
                # Tenta anexar ao ultimo documento assinavel (so para judicial etc)
                if resultados_processados:
                    ultimo = resultados_processados[-1]
                    cat_ultimo = classificar_doc(ultimo, tipo_processo)
                    if cat_ultimo in CATEGORIAS_ASSINADAS:
                        # Merge protocolo com o documento anterior
                        merged_path = os.path.join(tmp_dir, f"merged_{len(resultados_processados)}.pdf")
                        if merge_pdfs([ultimo["arquivo_tmp"], r["arquivo_tmp"]], merged_path):
                            ultimo["arquivo_tmp"] = merged_path
                            ultimo["nome_original"] += " + protocolo"
                            continue
                # Se nao conseguiu anexar, guarda para tentar depois
                protocolos_pendentes.append(r)
            else:
                resultados_processados.append(r)

        # ===== ETAPA 2: Classificar documentos =====
        sequencia = SEQUENCIA_JUDICIAL if tipo_processo != "inss_admin" else SEQUENCIA_INSS_ADMIN
        categorias_validas = [cat for cat, _ in sequencia]

        # Para admin, tambem inclui categorias individuais que serao reagrupadas
        if tipo_processo == "inss_admin":
            for grupo, componentes in GRUPO_MERGE_ADMIN.items():
                categorias_validas.extend(componentes)

        docs_por_categoria = {cat: [] for cat in categorias_validas}
        docs_cronologicos = []

        for r in resultados_processados:
            cat = classificar_doc(r, tipo_processo)
            if cat and cat in docs_por_categoria:
                docs_por_categoria[cat].append(r)
            else:
                docs_cronologicos.append(r)

        # ===== ETAPA 2.5: Para admin, agrupar procuracao+substabelecimento+termo =====
        if tipo_processo == "inss_admin":
            for grupo, componentes in GRUPO_MERGE_ADMIN.items():
                grupo_docs = []
                for comp in componentes:
                    grupo_docs.extend(docs_por_categoria.pop(comp, []))
                if grupo_docs:
                    # Ordena na sequencia: procuracao primeiro, depois substabelecimento, depois termo
                    ordem_componentes = {c: i for i, c in enumerate(componentes)}
                    grupo_docs.sort(key=lambda r: ordem_componentes.get(
                        classificar_doc(r, tipo_processo), 99))
                    docs_por_categoria[grupo] = grupo_docs

        # Ordena cronologicamente os demais
        docs_cronologicos.sort(key=lambda r: r.get("data") or "9999-99-99")

        # Ordena cronologicamente dentro das categorias MERGE (CTPS, atestados, etc.)
        for cat in CATEGORIAS_MERGE:
            if cat in docs_por_categoria and cat not in GRUPO_MERGE_ADMIN:
                docs_por_categoria[cat].sort(key=lambda r: r.get("data") or "9999-99-99")

        # Monta ZIP
        zip_buffer = io.BytesIO()
        nome_limpo = limpar_nome(nome_cliente)
        # UUID curto (8 chars) para evitar race condition entre usuarios simultaneos
        uid = uuid.uuid4().hex[:8]
        nome_pasta = f"{nome_limpo}_{tipo_processo}_{uid}"
        # Nome "amigavel" para a pasta dentro do ZIP (sem o UUID, mais limpo)
        nome_pasta_zip = f"{nome_limpo}_{tipo_processo}"
        limite_mb = LIMITES_TAMANHO.get(tipo_processo, MAX_FILE_SIZE_BYTES) // (1024 * 1024)
        relatorio_linhas = [
            f"Relatorio de Organizacao - {nome_cliente}",
            f"Tipo: {TIPOS_PROCESSO[tipo_processo]}",
            f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            f"Total de documentos separados: {len(resultados_processados)}",
            f"Limite por arquivo: {limite_mb}MB",
            "-" * 50,
            "",
        ]

        # Adiciona avisos no topo do relatorio (ex: PDFs truncados)
        avisos_relatorio = list({r.get("_aviso") for r in resultados if r.get("_aviso")})
        if avisos_relatorio:
            relatorio_linhas.insert(5, "")
            relatorio_linhas.insert(5, "AVISOS:")
            for a in avisos_relatorio:
                relatorio_linhas.insert(6, f"  - {a}")

        lista_docs = []
        ordem = 1
        limite_arquivo = LIMITES_TAMANHO.get(tipo_processo, MAX_FILE_SIZE_BYTES)

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            # ===== ETAPA 3: Adicionar documentos na sequencia correta =====
            for cat, label in sequencia:
                docs_cat = docs_por_categoria.get(cat, [])
                if not docs_cat:
                    continue

                # Se categoria e MERGE, junta todos em um unico PDF
                if cat in CATEGORIAS_MERGE and len(docs_cat) > 1:
                    merged_path = os.path.join(tmp_dir, f"merged_{cat}.pdf")
                    caminhos = [r["arquivo_tmp"] for r in docs_cat]
                    if merge_pdfs(caminhos, merged_path):
                        nome_label = label or limpar_nome(docs_cat[0].get("tipo", "documento"))[:30].upper()
                        partes = dividir_pdf_por_tamanho(merged_path, tmp_dir, max_bytes=limite_arquivo)
                        for idx, parte_path in enumerate(partes):
                            sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                            novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{sufixo}.pdf"
                            zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                            relatorio_linhas.append(
                                f"{novo_nome} | {len(docs_cat)} docs merged em ordem cronologica"
                            )
                            lista_docs.append({
                                "ordem": ordem, "nome": novo_nome,
                                "original": f"{len(docs_cat)} documentos merged",
                                "tipo": label or cat, "data": None,
                            })
                            ordem += 1
                        continue

                # Sem merge: adiciona cada doc individualmente
                for r in docs_cat:
                    nome_label = label or limpar_nome(r.get("tipo", "documento"))[:30].upper()
                    # Verifica tamanho e divide se necessario
                    if r["extensao"] == ".pdf" and os.path.getsize(r["arquivo_tmp"]) > limite_arquivo:
                        partes = dividir_pdf_por_tamanho(r["arquivo_tmp"], tmp_dir, max_bytes=limite_arquivo)
                        for idx, parte_path in enumerate(partes):
                            sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                            novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{sufixo}{r['extensao']}"
                            zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                            relatorio_linhas.append(
                                f"{novo_nome} | Original: {r['nome_original']} | Data: {r.get('data', 'N/A')}"
                            )
                            lista_docs.append({
                                "ordem": ordem, "nome": novo_nome,
                                "original": r["nome_original"],
                                "tipo": r.get("tipo", "?"), "data": r.get("data"),
                            })
                            ordem += 1
                    else:
                        novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{r['extensao']}"
                        zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                        relatorio_linhas.append(
                            f"{novo_nome} | Original: {r['nome_original']} | Tipo: {r.get('tipo', '?')} | Data: {r.get('data', 'N/A')}"
                        )
                        lista_docs.append({
                            "ordem": ordem, "nome": novo_nome,
                            "original": r["nome_original"],
                            "tipo": r.get("tipo", "?"), "data": r.get("data"),
                        })
                        ordem += 1

            # ===== ETAPA 4: Demais documentos em ordem cronologica =====
            for r in docs_cronologicos:
                tipo_limpo = limpar_nome(r.get("tipo", "documento"))[:30]
                data_str = r.get("data") or "sem_data"
                if r["extensao"] == ".pdf" and os.path.getsize(r["arquivo_tmp"]) > limite_arquivo:
                    partes = dividir_pdf_por_tamanho(r["arquivo_tmp"], tmp_dir, max_bytes=limite_arquivo)
                    for idx, parte_path in enumerate(partes):
                        sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                        novo_nome = f"{ordem:02d}_{nome_limpo}_{data_str}_{tipo_limpo}{sufixo}{r['extensao']}"
                        zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                        fonte = r.get("_data_fonte")
                        fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                        relatorio_linhas.append(
                            f"{novo_nome} | Original: {r['nome_original']} | Data: {r.get('data', 'NAO ENCONTRADA')}{fonte_txt}"
                        )
                        lista_docs.append({
                            "ordem": ordem, "nome": novo_nome,
                            "original": r["nome_original"],
                            "tipo": r.get("tipo", "?"), "data": r.get("data"),
                            "data_fonte": r.get("_data_fonte"),
                        })
                        ordem += 1
                else:
                    novo_nome = f"{ordem:02d}_{nome_limpo}_{data_str}_{tipo_limpo}{r['extensao']}"
                    zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                    fonte = r.get("_data_fonte")
                    fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                    relatorio_linhas.append(
                        f"{novo_nome} | Original: {r['nome_original']} | Tipo: {r.get('tipo', '?')} | Data: {r.get('data', 'NAO ENCONTRADA')}{fonte_txt}"
                    )
                    lista_docs.append({
                        "ordem": ordem, "nome": novo_nome,
                        "original": r["nome_original"],
                        "tipo": r.get("tipo", "?"), "data": r.get("data"),
                        "data_fonte": r.get("_data_fonte"),
                    })
                    ordem += 1

            # Protocolos orfaos (caso nao tenham sido anexados)
            for r in protocolos_pendentes:
                novo_nome = f"{ordem:02d}_{nome_limpo}_Protocolo_Assinatura{r['extensao']}"
                zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                relatorio_linhas.append(
                    f"{novo_nome} | Original: {r['nome_original']} | PROTOCOLO ORFAO"
                )
                lista_docs.append({
                    "ordem": ordem, "nome": novo_nome,
                    "original": r["nome_original"],
                    "tipo": "Protocolo Assinatura", "data": None,
                })
                ordem += 1

            relatorio_linhas.append("")
            relatorio_linhas.append("Gerado por: Organizador Juridico AB Group")
            zf.writestr(f"{nome_pasta_zip}/_relatorio.txt", "\n".join(relatorio_linhas))

        zip_buffer.seek(0)

        shared_zip = SHARED_ZIP_DIR / f"{nome_pasta}.zip"
        with open(shared_zip, "wb") as f:
            f.write(zip_buffer.getvalue())

        shutil.rmtree(tmp_dir, ignore_errors=True)

        # Registra o uso na planilha (via webhook n8n)
        registrar_uso(usuario, nome_cliente, tipo_processo, len(resultados), "sucesso")

        # Coleta avisos (ex: PDF truncado por ter mais que MAX_PAGINAS_PDF)
        avisos = list({r.get("_aviso") for r in resultados if r.get("_aviso")})

        # Registra avisos como eventos de erro menores (pra monitoramento)
        for aviso in avisos:
            registrar_erro(usuario, nome_cliente, tipo_processo, "",
                          "pdf_truncado", aviso)

        # Registra se muitos documentos sairam como desconhecido (qualidade ruim)
        total_desconhecidos = sum(1 for r in resultados if r.get("tipo", "").lower() == "desconhecido")
        if total_desconhecidos > 0 and total_desconhecidos >= len(resultados) * 0.3:
            registrar_erro(usuario, nome_cliente, tipo_processo, "",
                          "muitos_desconhecidos",
                          f"{total_desconhecidos} de {len(resultados)} documentos nao foram classificados")

        return jsonify({
            "sucesso": True,
            "nome_pasta": nome_pasta,
            "documentos": lista_docs,
            "total": len(resultados),
            "com_data": sum(1 for r in resultados if r.get("data")),
            "sem_data": sum(1 for r in resultados if not r.get("data")),
            "avisos": avisos,
        })

    except Exception as e:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        registrar_uso(usuario, nome_cliente, tipo_processo, 0, f"erro: {str(e)[:100]}")
        registrar_erro(usuario, nome_cliente, tipo_processo, "",
                      "excecao_geral_processar", e)
        return jsonify({"erro": str(e)}), 500


def _sse(evento):
    """Formata um dicionario como evento SSE."""
    return f"data: {json.dumps(evento)}\n\n"


@app.route("/processar-stream", methods=["POST"])
def processar_stream():
    """Versao streaming do /processar. Envia eventos SSE de progresso em tempo real.
    Eventos:
      - {"tipo": "progresso", "etapa": "...", "percent": 0-100}
      - {"tipo": "complete", "resultado": {...mesma payload do /processar}}
      - {"tipo": "error", "mensagem": "..."}
    """
    import anthropic

    limpar_zips_antigos()
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return jsonify({"erro": "API key nao configurada no servidor"}), 500

    usuario = request.form.get("usuario", "").strip()
    nome_cliente = request.form.get("nome_cliente", "").strip()
    tipo_processo = request.form.get("tipo_processo", "").strip()

    if not usuario or usuario not in USUARIOS:
        return jsonify({"erro": "Selecione o usuario que esta processando"}), 400
    if not nome_cliente:
        return jsonify({"erro": "Nome do cliente e obrigatorio"}), 400
    if tipo_processo not in TIPOS_PROCESSO:
        return jsonify({"erro": "Tipo de processo invalido"}), 400

    arquivos = request.files.getlist("documentos")
    if not arquivos or all(f.filename == "" for f in arquivos):
        return jsonify({"erro": "Nenhum documento enviado"}), 400

    arquivos_validos = [
        f for f in arquivos
        if f.filename and Path(f.filename).suffix.lower() in EXTENSOES_ACEITAS
    ]
    if not arquivos_validos:
        return jsonify({"erro": f"Nenhum arquivo valido. Aceitos: {', '.join(EXTENSOES_ACEITAS)}"}), 400

    tmp_dir = tempfile.mkdtemp()

    # Salva arquivos em disco antes de iniciar stream (request nao da pra ler dentro do generator)
    arquivos_para_processar = []
    for arquivo in arquivos_validos:
        nome_original = arquivo.filename
        ext = Path(nome_original).suffix.lower()
        tmp_path = os.path.join(tmp_dir, nome_original)
        arquivo.save(tmp_path)
        arquivos_para_processar.append((nome_original, tmp_path, ext))

    # Log de inicio de processamento (auditoria)
    audit_log.info(
        "INICIO | usuario=%s | cliente=%s | tipo=%s | arquivos=%d | nomes=%s",
        usuario, nome_cliente, tipo_processo,
        len(arquivos_para_processar),
        ",".join(a[0] for a in arquivos_para_processar)[:500],
    )
    # Persiste no SQLite local
    audit_save(
        evento="INICIO",
        usuario=usuario,
        cliente=nome_cliente,
        tipo_processo=TIPOS_PROCESSO.get(tipo_processo, tipo_processo),
        arquivos=[a[0] for a in arquivos_para_processar],
        total_documentos=len(arquivos_para_processar),
    )

    @stream_with_context
    def gerar():
        try:
            yield _sse({"tipo": "progresso", "etapa": "Iniciando processamento...", "percent": 2})

            client = anthropic.Anthropic(api_key=api_key)
            resultados = []
            total = len(arquivos_para_processar)

            # ===== ETAPA: ANALISE DE CADA ARQUIVO =====
            for i, (nome_original, tmp_path, ext) in enumerate(arquivos_para_processar):
                pct = 5 + int((i / total) * 55)  # 5% a 60%
                yield _sse({
                    "tipo": "progresso",
                    "etapa": f"Analisando arquivo {i+1} de {total}: {nome_original}",
                    "percent": pct,
                })
                try:
                    docs_separados = processar_arquivo_completo(client, tmp_path, nome_original, tmp_dir)
                    for doc in docs_separados:
                        resolver_data_fallback(doc, tmp_path, nome_original)
                    resultados.extend(docs_separados)
                except Exception as e:
                    resultados.append({
                        "tipo": "desconhecido", "data": None,
                        "arquivo_tmp": tmp_path, "nome_original": nome_original,
                        "extensao": ext, "_data_fonte": None,
                    })
                    registrar_erro(usuario, nome_cliente, tipo_processo, nome_original,
                                  "excecao_processamento_arquivo", e)

            # ===== ETAPA: DUPLA CHECAGEM DE DOCS CRITICOS (#3.2) =====
            # 2a chamada da IA pra confirmar classificacao de procuracoes,
            # declaracoes, contratos, decisoes INSS — tipos onde erro tem custo alto
            criticos = [r for r in resultados if eh_tipo_critico(r.get("tipo", ""))]
            if criticos:
                yield _sse({
                    "tipo": "progresso",
                    "etapa": f"Validando {len(criticos)} documento(s) critico(s)...",
                    "percent": 62,
                })
                for idx, doc in enumerate(criticos):
                    try:
                        dupla_checagem_doc(client, doc)
                    except Exception:
                        pass

            # ===== ETAPA: DETECTAR DUPLICATAS (#3.4) =====
            yield _sse({"tipo": "progresso", "etapa": "Detectando duplicatas...", "percent": 64})
            num_dup = detectar_duplicatas(resultados)
            if num_dup > 0:
                # Filtra duplicatas do fluxo direto (nao vao pro ZIP montado automaticamente)
                # Mas continuam disponiveis na sessao pra usuario ver na revisao
                resultados_sem_dup = [r for r in resultados if not r.get("_duplicata")]
            else:
                resultados_sem_dup = resultados
            # IMPORTANTE: usamos resultados_sem_dup pro ZIP "direto", e resultados completos pra sessao

            # ===== ETAPA 1: Anexar protocolos =====
            yield _sse({"tipo": "progresso", "etapa": "Agrupando protocolos de assinatura...", "percent": 65})
            resultados_processados = []
            protocolos_pendentes = []
            for r in resultados_sem_dup:
                if eh_protocolo_assinatura(r) and tipo_processo != "inss_admin":
                    if resultados_processados:
                        ultimo = resultados_processados[-1]
                        cat_ultimo = classificar_doc(ultimo, tipo_processo)
                        if cat_ultimo in CATEGORIAS_ASSINADAS:
                            merged_path = os.path.join(tmp_dir, f"merged_{len(resultados_processados)}.pdf")
                            if merge_pdfs([ultimo["arquivo_tmp"], r["arquivo_tmp"]], merged_path):
                                ultimo["arquivo_tmp"] = merged_path
                                ultimo["nome_original"] += " + protocolo"
                                continue
                    protocolos_pendentes.append(r)
                else:
                    resultados_processados.append(r)

            # ===== ETAPA 2: Classificar =====
            yield _sse({"tipo": "progresso", "etapa": "Classificando documentos...", "percent": 75})
            sequencia = SEQUENCIA_JUDICIAL if tipo_processo != "inss_admin" else SEQUENCIA_INSS_ADMIN
            categorias_validas = [cat for cat, _ in sequencia]
            if tipo_processo == "inss_admin":
                for grupo, componentes in GRUPO_MERGE_ADMIN.items():
                    categorias_validas.extend(componentes)
            docs_por_categoria = {cat: [] for cat in categorias_validas}
            docs_cronologicos = []
            for r in resultados_processados:
                cat = classificar_doc(r, tipo_processo)
                if cat and cat in docs_por_categoria:
                    docs_por_categoria[cat].append(r)
                else:
                    docs_cronologicos.append(r)

            # ETAPA 2.5: admin grouping
            if tipo_processo == "inss_admin":
                for grupo, componentes in GRUPO_MERGE_ADMIN.items():
                    grupo_docs = []
                    for comp in componentes:
                        grupo_docs.extend(docs_por_categoria.pop(comp, []))
                    if grupo_docs:
                        ordem_componentes = {c: i for i, c in enumerate(componentes)}
                        grupo_docs.sort(key=lambda r: ordem_componentes.get(
                            classificar_doc(r, tipo_processo), 99))
                        docs_por_categoria[grupo] = grupo_docs

            docs_cronologicos.sort(key=lambda r: r.get("data") or "9999-99-99")
            for cat in CATEGORIAS_MERGE:
                if cat in docs_por_categoria and cat not in GRUPO_MERGE_ADMIN:
                    docs_por_categoria[cat].sort(key=lambda r: r.get("data") or "9999-99-99")

            # ===== ETAPA 3: Montar ZIP =====
            yield _sse({"tipo": "progresso", "etapa": "Gerando arquivo ZIP...", "percent": 85})
            zip_buffer = io.BytesIO()
            nome_limpo = limpar_nome(nome_cliente)
            uid = uuid.uuid4().hex[:8]
            nome_pasta = f"{nome_limpo}_{tipo_processo}_{uid}"
            nome_pasta_zip = f"{nome_limpo}_{tipo_processo}"
            limite_mb = LIMITES_TAMANHO.get(tipo_processo, MAX_FILE_SIZE_BYTES) // (1024 * 1024)
            relatorio_linhas = [
                f"Relatorio de Organizacao - {nome_cliente}",
                f"Tipo: {TIPOS_PROCESSO[tipo_processo]}",
                f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                f"Total de documentos separados: {len(resultados_processados)}",
                f"Limite por arquivo: {limite_mb}MB",
                "-" * 50, "",
            ]
            avisos_relatorio = list({r.get("_aviso") for r in resultados if r.get("_aviso")})
            if avisos_relatorio:
                relatorio_linhas.insert(5, "")
                relatorio_linhas.insert(5, "AVISOS:")
                for a in avisos_relatorio:
                    relatorio_linhas.insert(6, f"  - {a}")

            lista_docs = []
            ordem = 1
            limite_arquivo = LIMITES_TAMANHO.get(tipo_processo, MAX_FILE_SIZE_BYTES)

            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                for cat, label in sequencia:
                    docs_cat = docs_por_categoria.get(cat, [])
                    if not docs_cat:
                        continue
                    if cat in CATEGORIAS_MERGE and len(docs_cat) > 1:
                        merged_path = os.path.join(tmp_dir, f"merged_{cat}.pdf")
                        caminhos = [r["arquivo_tmp"] for r in docs_cat]
                        if merge_pdfs(caminhos, merged_path):
                            nome_label = label or limpar_nome(docs_cat[0].get("tipo", "documento"))[:30].upper()
                            partes = dividir_pdf_por_tamanho(merged_path, tmp_dir, max_bytes=limite_arquivo)
                            for idx, parte_path in enumerate(partes):
                                sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                                novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{sufixo}.pdf"
                                zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                                relatorio_linhas.append(f"{novo_nome} | {len(docs_cat)} docs merged")
                                lista_docs.append({
                                    "ordem": ordem, "nome": novo_nome,
                                    "original": f"{len(docs_cat)} documentos merged",
                                    "tipo": label or cat, "data": None,
                                })
                                ordem += 1
                            continue
                    for r in docs_cat:
                        nome_label = label or limpar_nome(r.get("tipo", "documento"))[:30].upper()
                        if r["extensao"] == ".pdf" and os.path.getsize(r["arquivo_tmp"]) > limite_arquivo:
                            partes = dividir_pdf_por_tamanho(r["arquivo_tmp"], tmp_dir, max_bytes=limite_arquivo)
                            for idx, parte_path in enumerate(partes):
                                sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                                novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{sufixo}{r['extensao']}"
                                zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                                fonte = r.get("_data_fonte")
                                fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                                relatorio_linhas.append(
                                    f"{novo_nome} | Original: {r['nome_original']} | Data: {r.get('data', 'N/A')}{fonte_txt}"
                                )
                                lista_docs.append({
                                    "ordem": ordem, "nome": novo_nome,
                                    "original": r["nome_original"],
                                    "tipo": r.get("tipo", "?"), "data": r.get("data"),
                                    "data_fonte": r.get("_data_fonte"),
                                })
                                ordem += 1
                        else:
                            novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{r['extensao']}"
                            zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                            fonte = r.get("_data_fonte")
                            fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                            relatorio_linhas.append(
                                f"{novo_nome} | Original: {r['nome_original']} | Tipo: {r.get('tipo', '?')} | Data: {r.get('data', 'N/A')}{fonte_txt}"
                            )
                            lista_docs.append({
                                "ordem": ordem, "nome": novo_nome,
                                "original": r["nome_original"],
                                "tipo": r.get("tipo", "?"), "data": r.get("data"),
                                "data_fonte": r.get("_data_fonte"),
                            })
                            ordem += 1

                for r in docs_cronologicos:
                    tipo_limpo = limpar_nome(r.get("tipo", "documento"))[:30]
                    data_str = r.get("data") or "sem_data"
                    if r["extensao"] == ".pdf" and os.path.getsize(r["arquivo_tmp"]) > limite_arquivo:
                        partes = dividir_pdf_por_tamanho(r["arquivo_tmp"], tmp_dir, max_bytes=limite_arquivo)
                        for idx, parte_path in enumerate(partes):
                            sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                            novo_nome = f"{ordem:02d}_{nome_limpo}_{data_str}_{tipo_limpo}{sufixo}{r['extensao']}"
                            zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                            fonte = r.get("_data_fonte")
                            fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                            relatorio_linhas.append(
                                f"{novo_nome} | Original: {r['nome_original']} | Data: {r.get('data', 'NAO ENCONTRADA')}{fonte_txt}"
                            )
                            lista_docs.append({
                                "ordem": ordem, "nome": novo_nome,
                                "original": r["nome_original"],
                                "tipo": r.get("tipo", "?"), "data": r.get("data"),
                                "data_fonte": r.get("_data_fonte"),
                            })
                            ordem += 1
                    else:
                        novo_nome = f"{ordem:02d}_{nome_limpo}_{data_str}_{tipo_limpo}{r['extensao']}"
                        zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                        fonte = r.get("_data_fonte")
                        fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                        relatorio_linhas.append(
                            f"{novo_nome} | Original: {r['nome_original']} | Tipo: {r.get('tipo', '?')} | Data: {r.get('data', 'NAO ENCONTRADA')}{fonte_txt}"
                        )
                        lista_docs.append({
                            "ordem": ordem, "nome": novo_nome,
                            "original": r["nome_original"],
                            "tipo": r.get("tipo", "?"), "data": r.get("data"),
                            "data_fonte": r.get("_data_fonte"),
                        })
                        ordem += 1

                for r in protocolos_pendentes:
                    novo_nome = f"{ordem:02d}_{nome_limpo}_Protocolo_Assinatura{r['extensao']}"
                    zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                    relatorio_linhas.append(f"{novo_nome} | PROTOCOLO ORFAO")
                    lista_docs.append({
                        "ordem": ordem, "nome": novo_nome,
                        "original": r["nome_original"],
                        "tipo": "Protocolo Assinatura", "data": None,
                    })
                    ordem += 1

                relatorio_linhas.append("")
                relatorio_linhas.append("Gerado por: Organizador Juridico AB Group")
                zf.writestr(f"{nome_pasta_zip}/_relatorio.txt", "\n".join(relatorio_linhas))

            # ===== ETAPA 4: Salvar e registrar =====
            yield _sse({"tipo": "progresso", "etapa": "Finalizando...", "percent": 95})
            zip_buffer.seek(0)
            shared_zip = SHARED_ZIP_DIR / f"{nome_pasta}.zip"
            with open(shared_zip, "wb") as f:
                f.write(zip_buffer.getvalue())

            # Salva sessao de revisao (#2.2): usuario pode editar tipos/datas e gerar ZIP novamente
            limpar_sessoes_antigas()
            session_id = uuid.uuid4().hex
            # docs_para_sessao = todos os docs antes do merge (flat) pra permitir edicao individual
            docs_para_sessao = list(resultados_processados) + list(protocolos_pendentes)
            contexto_sessao = {
                "usuario": usuario,
                "nome_cliente": nome_cliente,
                "tipo_processo": tipo_processo,
            }
            try:
                salvar_sessao(session_id, contexto_sessao, docs_para_sessao)
            except Exception:
                session_id = None  # se falhar, frontend cai no fluxo direto

            shutil.rmtree(tmp_dir, ignore_errors=True)

            registrar_uso(usuario, nome_cliente, tipo_processo, len(resultados), "sucesso")
            avisos = list({r.get("_aviso") for r in resultados if r.get("_aviso")})
            for aviso in avisos:
                registrar_erro(usuario, nome_cliente, tipo_processo, "", "pdf_truncado", aviso)
            total_desconhecidos = sum(1 for r in resultados if r.get("tipo", "").lower() == "desconhecido")
            if total_desconhecidos > 0 and total_desconhecidos >= len(resultados) * 0.3:
                registrar_erro(usuario, nome_cliente, tipo_processo, "", "muitos_desconhecidos",
                              f"{total_desconhecidos} de {len(resultados)} documentos nao classificados")

            # Preparar lista de docs brutos pra revisao (usando dados da sessao)
            docs_revisao = []
            if session_id:
                sess = ler_sessao(session_id)
                if sess:
                    docs_revisao = sess["docs"]

            yield _sse({
                "tipo": "complete",
                "resultado": {
                    "sucesso": True,
                    "nome_pasta": nome_pasta,
                    "session_id": session_id,
                    "documentos": lista_docs,
                    "docs_revisao": docs_revisao,  # lista flat pra UI de edicao
                    "total": len(resultados),
                    "com_data": sum(1 for r in resultados if r.get("data")),
                    "sem_data": sum(1 for r in resultados if not r.get("data")),
                    "avisos": avisos,
                }
            })

        except Exception as e:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            registrar_uso(usuario, nome_cliente, tipo_processo, 0, f"erro: {str(e)[:100]}")
            registrar_erro(usuario, nome_cliente, tipo_processo, "", "excecao_geral_processar", e)
            yield _sse({"tipo": "error", "mensagem": str(e)})

    return Response(gerar(), mimetype='text/event-stream', headers={
        'X-Accel-Buffering': 'no',
        'Cache-Control': 'no-cache',
    })


@app.route("/gerar-zip-revisado/<session_id>", methods=["POST"])
def gerar_zip_revisado(session_id):
    """Recebe lista de docs editados pelo usuario e monta o ZIP final.
    Body JSON: {"docs": [{id, tipo, data, deletar, ordem?}, ...]}
    """
    sess = ler_sessao(session_id)
    if not sess:
        return jsonify({"erro": "Sessao expirada ou nao encontrada. Processe novamente."}), 404

    try:
        body = request.get_json(silent=True) or {}
        docs_edit_raw = body.get("docs", [])

        # Mapa dos docs originais por id pra preencher campos nao enviados
        docs_orig_map = {d["id"]: d for d in sess["docs"]}

        # Junta edicoes com dados originais, preservando a ordem enviada pelo frontend
        docs_editados = []
        for d_edit in docs_edit_raw:
            orig = docs_orig_map.get(d_edit.get("id"))
            if not orig:
                continue
            # Merge: usa valores editados, ou os originais se nao enviou
            docs_editados.append({
                "id": orig["id"],
                "arquivo": orig["arquivo"],
                "extensao": orig["extensao"],
                "nome_original": orig.get("nome_original", ""),
                "tipo": d_edit.get("tipo", orig.get("tipo")),
                "data": d_edit.get("data", orig.get("data")),
                "data_fonte": orig.get("data_fonte"),
                "aviso": orig.get("aviso"),
                "deletar": bool(d_edit.get("deletar", False)),
            })

        # Monta ZIP num tmp_dir novo
        tmp_dir = tempfile.mkdtemp()
        try:
            resultado = montar_zip_final(sess["contexto"], docs_editados, sess["sess_dir"], tmp_dir)
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        contexto = sess["contexto"]
        registrar_uso(contexto.get("usuario"), contexto.get("nome_cliente"),
                     contexto.get("tipo_processo"), resultado["total"], "sucesso_revisado")

        return jsonify({
            "sucesso": True,
            "nome_pasta": resultado["nome_pasta"],
            "documentos": resultado["lista_docs"],
            "total": resultado["total"],
            "com_data": resultado["com_data"],
            "sem_data": resultado["sem_data"],
            "avisos": resultado["avisos"],
        })

    except Exception as e:
        registrar_erro(sess["contexto"].get("usuario"), sess["contexto"].get("nome_cliente"),
                      sess["contexto"].get("tipo_processo"), "",
                      "excecao_gerar_zip_revisado", e)
        return jsonify({"erro": str(e)}), 500


@app.route("/download/<nome_pasta>")
def download(nome_pasta):
    zip_path = SHARED_ZIP_DIR / f"{nome_pasta}.zip"
    if not zip_path.exists():
        return "Arquivo nao encontrado. Processe novamente.", 404

    # Remove UUID do nome do arquivo baixado (fica amigavel pro usuario)
    # nome_pasta = "joao_silva_inss_admin_abc12345" -> download "joao_silva_inss_admin_organizado.zip"
    nome_amigavel = nome_pasta.rsplit("_", 1)[0] if "_" in nome_pasta else nome_pasta
    return send_file(
        zip_path,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"{nome_amigavel}_organizado.zip",
    )


def limpar_zips_antigos(max_idade_segundos=3600):
    """Remove ZIPs com mais de 1 hora do diretorio compartilhado."""
    try:
        agora = time.time()
        for zip_file in SHARED_ZIP_DIR.glob("*.zip"):
            if agora - zip_file.stat().st_mtime > max_idade_segundos:
                try:
                    zip_file.unlink()
                except Exception:
                    pass
    except Exception:
        pass


def limpar_sessoes_antigas(max_idade_segundos=3600):
    """Remove diretorios de sessao com mais de 1 hora."""
    try:
        agora = time.time()
        for sess_dir in SESSIONS_DIR.iterdir():
            if sess_dir.is_dir() and agora - sess_dir.stat().st_mtime > max_idade_segundos:
                try:
                    shutil.rmtree(sess_dir, ignore_errors=True)
                except Exception:
                    pass
    except Exception:
        pass


def salvar_sessao(session_id, contexto, docs):
    """Salva uma sessao de revisao em disco.

    contexto: dict com usuario, nome_cliente, tipo_processo, limite_arquivo, nome_limpo, nome_pasta_zip
    docs: lista flat de dicts com arquivo_tmp, tipo, data, extensao, nome_original, _data_fonte, _aviso
          (os arquivos sao copiados para SESSIONS_DIR/<session_id>/)
    """
    sess_dir = SESSIONS_DIR / session_id
    sess_dir.mkdir(exist_ok=True)

    docs_meta = []
    for i, doc in enumerate(docs):
        arq_src = doc.get("arquivo_tmp")
        if not arq_src or not os.path.exists(arq_src):
            continue
        ext = doc.get("extensao") or Path(arq_src).suffix.lower() or ".pdf"
        nome_arq_sessao = f"doc-{i:03d}{ext}"
        dest = sess_dir / nome_arq_sessao
        try:
            shutil.copy2(arq_src, dest)
        except Exception:
            continue
        docs_meta.append({
            "id": f"doc-{i:03d}",
            "arquivo": nome_arq_sessao,
            "tipo": doc.get("tipo", "desconhecido"),
            "data": doc.get("data"),
            "data_fonte": doc.get("_data_fonte"),
            "nome_original": doc.get("nome_original", ""),
            "aviso": doc.get("_aviso"),
            "extensao": ext,
            # Dupla checagem (#3.2)
            "confianca": doc.get("_confianca"),  # "alta" | "media" | None
            "tipo_alternativo": doc.get("_tipo_alternativo"),
            "razao_dupla_checagem": doc.get("_razao_dupla_checagem"),
            # Duplicatas (#3.4)
            "duplicata": doc.get("_duplicata", False),
            "duplicata_de": doc.get("_duplicata_de"),
            "motivo_duplicata": doc.get("_motivo_duplicata"),
        })

    metadata = {
        "contexto": contexto,
        "docs": docs_meta,
        "created_at": datetime.now().isoformat(),
    }
    with open(sess_dir / "metadata.json", "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)
    return metadata


def ler_sessao(session_id):
    """Le metadata e retorna contexto + docs. None se sessao nao existe."""
    sess_dir = SESSIONS_DIR / session_id
    meta_path = sess_dir / "metadata.json"
    if not meta_path.exists():
        return None
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            meta = json.load(f)
        meta["sess_dir"] = str(sess_dir)
        return meta
    except Exception:
        return None


def montar_zip_final(contexto, docs_editados, sess_dir, tmp_dir):
    """Monta o ZIP final a partir de uma lista de docs (ja editados pelo usuario).

    docs_editados: lista de dicts [{id, tipo, data, deletar, arquivo, extensao, nome_original, aviso, data_fonte}]
      Nota: ordem da lista = ordem final no ZIP (menos merges de grupo/categoria).
    Aplica classificacao e merge de grupos baseado nos TIPOS editados.
    Retorna dict com nome_pasta, lista_docs, relatorio, zip_bytes.
    """
    tipo_processo = contexto["tipo_processo"]
    nome_cliente = contexto["nome_cliente"]
    usuario = contexto.get("usuario", "")
    limite_arquivo = LIMITES_TAMANHO.get(tipo_processo, MAX_FILE_SIZE_BYTES)

    # Filtra docs deletados
    docs_ativos = [d for d in docs_editados if not d.get("deletar")]

    # Recria a estrutura de arquivo_tmp apontando pros arquivos da sessao
    for d in docs_ativos:
        d["arquivo_tmp"] = os.path.join(sess_dir, d["arquivo"])

    # Aplica mesma logica de ETAPA 1 (protocolos) + ETAPA 2 (classificacao) + ETAPA 2.5 (grupos)
    # Mas preservando a ordem que o usuario deixou na lista

    # ETAPA 1 — protocolos em processos judiciais: anexa ao doc anterior
    resultados_processados = []
    protocolos_pendentes = []
    for r in docs_ativos:
        if eh_protocolo_assinatura(r) and tipo_processo != "inss_admin":
            if resultados_processados:
                ultimo = resultados_processados[-1]
                cat_ultimo = classificar_doc(ultimo, tipo_processo)
                if cat_ultimo in CATEGORIAS_ASSINADAS:
                    merged_path = os.path.join(tmp_dir, f"merged_{len(resultados_processados)}.pdf")
                    if merge_pdfs([ultimo["arquivo_tmp"], r["arquivo_tmp"]], merged_path):
                        ultimo["arquivo_tmp"] = merged_path
                        ultimo["nome_original"] = (ultimo.get("nome_original") or "") + " + protocolo"
                        continue
            protocolos_pendentes.append(r)
        else:
            resultados_processados.append(r)

    # ETAPA 2 — classificar
    sequencia = SEQUENCIA_JUDICIAL if tipo_processo != "inss_admin" else SEQUENCIA_INSS_ADMIN
    categorias_validas = [cat for cat, _ in sequencia]
    if tipo_processo == "inss_admin":
        for grupo, componentes in GRUPO_MERGE_ADMIN.items():
            categorias_validas.extend(componentes)
    docs_por_categoria = {cat: [] for cat in categorias_validas}
    docs_cronologicos = []
    for r in resultados_processados:
        cat = classificar_doc(r, tipo_processo)
        if cat and cat in docs_por_categoria:
            docs_por_categoria[cat].append(r)
        else:
            docs_cronologicos.append(r)

    # ETAPA 2.5 — agrupamento admin
    if tipo_processo == "inss_admin":
        for grupo, componentes in GRUPO_MERGE_ADMIN.items():
            grupo_docs = []
            for comp in componentes:
                grupo_docs.extend(docs_por_categoria.pop(comp, []))
            if grupo_docs:
                # Mantem a ordem que o usuario deixou
                docs_por_categoria[grupo] = grupo_docs

    # Ordena cronologicos por data
    docs_cronologicos.sort(key=lambda r: r.get("data") or "9999-99-99")
    for cat in CATEGORIAS_MERGE:
        if cat in docs_por_categoria and cat not in GRUPO_MERGE_ADMIN:
            docs_por_categoria[cat].sort(key=lambda r: r.get("data") or "9999-99-99")

    # ETAPA 3 — monta ZIP
    zip_buffer = io.BytesIO()
    nome_limpo = limpar_nome(nome_cliente)
    uid = uuid.uuid4().hex[:8]
    nome_pasta = f"{nome_limpo}_{tipo_processo}_{uid}"
    nome_pasta_zip = f"{nome_limpo}_{tipo_processo}"
    limite_mb = limite_arquivo // (1024 * 1024)
    relatorio_linhas = [
        f"Relatorio de Organizacao - {nome_cliente}",
        f"Tipo: {TIPOS_PROCESSO[tipo_processo]}",
        f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"Total de documentos separados: {len(resultados_processados)}",
        f"Limite por arquivo: {limite_mb}MB",
        "-" * 50, "",
    ]
    avisos_relatorio = list({r.get("aviso") for r in docs_ativos if r.get("aviso")})
    if avisos_relatorio:
        relatorio_linhas.insert(5, "")
        relatorio_linhas.insert(5, "AVISOS:")
        for a in avisos_relatorio:
            relatorio_linhas.insert(6, f"  - {a}")

    lista_docs = []
    ordem = 1

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for cat, label in sequencia:
            docs_cat = docs_por_categoria.get(cat, [])
            if not docs_cat:
                continue
            if cat in CATEGORIAS_MERGE and len(docs_cat) > 1:
                merged_path = os.path.join(tmp_dir, f"merged_{cat}.pdf")
                caminhos = [r["arquivo_tmp"] for r in docs_cat]
                if merge_pdfs(caminhos, merged_path):
                    nome_label = label or limpar_nome(docs_cat[0].get("tipo", "documento"))[:30].upper()
                    partes = dividir_pdf_por_tamanho(merged_path, tmp_dir, max_bytes=limite_arquivo)
                    for idx, parte_path in enumerate(partes):
                        sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                        novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{sufixo}.pdf"
                        zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                        relatorio_linhas.append(f"{novo_nome} | {len(docs_cat)} docs merged")
                        lista_docs.append({
                            "ordem": ordem, "nome": novo_nome,
                            "original": f"{len(docs_cat)} documentos merged",
                            "tipo": label or cat, "data": None,
                        })
                        ordem += 1
                    continue
            for r in docs_cat:
                nome_label = label or limpar_nome(r.get("tipo", "documento"))[:30].upper()
                ext = r.get("extensao", ".pdf")
                if ext == ".pdf" and os.path.getsize(r["arquivo_tmp"]) > limite_arquivo:
                    partes = dividir_pdf_por_tamanho(r["arquivo_tmp"], tmp_dir, max_bytes=limite_arquivo)
                    for idx, parte_path in enumerate(partes):
                        sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                        novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{sufixo}{ext}"
                        zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                        fonte = r.get("data_fonte")
                        fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                        relatorio_linhas.append(
                            f"{novo_nome} | Original: {r.get('nome_original','')} | Data: {r.get('data', 'N/A')}{fonte_txt}"
                        )
                        lista_docs.append({
                            "ordem": ordem, "nome": novo_nome,
                            "original": r.get("nome_original", ""),
                            "tipo": r.get("tipo", "?"), "data": r.get("data"),
                            "data_fonte": r.get("data_fonte"),
                        })
                        ordem += 1
                else:
                    novo_nome = f"{ordem:02d}_{nome_limpo}_{nome_label}{ext}"
                    zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                    fonte = r.get("data_fonte")
                    fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                    relatorio_linhas.append(
                        f"{novo_nome} | Original: {r.get('nome_original','')} | Tipo: {r.get('tipo','?')} | Data: {r.get('data', 'N/A')}{fonte_txt}"
                    )
                    lista_docs.append({
                        "ordem": ordem, "nome": novo_nome,
                        "original": r.get("nome_original", ""),
                        "tipo": r.get("tipo", "?"), "data": r.get("data"),
                        "data_fonte": r.get("data_fonte"),
                    })
                    ordem += 1

        for r in docs_cronologicos:
            tipo_limpo = limpar_nome(r.get("tipo", "documento"))[:30]
            data_str = r.get("data") or "sem_data"
            ext = r.get("extensao", ".pdf")
            if ext == ".pdf" and os.path.getsize(r["arquivo_tmp"]) > limite_arquivo:
                partes = dividir_pdf_por_tamanho(r["arquivo_tmp"], tmp_dir, max_bytes=limite_arquivo)
                for idx, parte_path in enumerate(partes):
                    sufixo = f"_parte{idx+1}" if len(partes) > 1 else ""
                    novo_nome = f"{ordem:02d}_{nome_limpo}_{data_str}_{tipo_limpo}{sufixo}{ext}"
                    zf.write(parte_path, f"{nome_pasta_zip}/{novo_nome}")
                    fonte = r.get("data_fonte")
                    fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                    relatorio_linhas.append(
                        f"{novo_nome} | Original: {r.get('nome_original','')} | Data: {r.get('data', 'NAO ENCONTRADA')}{fonte_txt}"
                    )
                    lista_docs.append({
                        "ordem": ordem, "nome": novo_nome,
                        "original": r.get("nome_original", ""),
                        "tipo": r.get("tipo", "?"), "data": r.get("data"),
                        "data_fonte": r.get("data_fonte"),
                    })
                    ordem += 1
            else:
                novo_nome = f"{ordem:02d}_{nome_limpo}_{data_str}_{tipo_limpo}{ext}"
                zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
                fonte = r.get("data_fonte")
                fonte_txt = f" [fonte: {fonte}]" if fonte and fonte != "ia" else ""
                relatorio_linhas.append(
                    f"{novo_nome} | Original: {r.get('nome_original','')} | Tipo: {r.get('tipo','?')} | Data: {r.get('data', 'NAO ENCONTRADA')}{fonte_txt}"
                )
                lista_docs.append({
                    "ordem": ordem, "nome": novo_nome,
                    "original": r.get("nome_original", ""),
                    "tipo": r.get("tipo", "?"), "data": r.get("data"),
                    "data_fonte": r.get("data_fonte"),
                })
                ordem += 1

        for r in protocolos_pendentes:
            ext = r.get("extensao", ".pdf")
            novo_nome = f"{ordem:02d}_{nome_limpo}_Protocolo_Assinatura{ext}"
            zf.write(r["arquivo_tmp"], f"{nome_pasta_zip}/{novo_nome}")
            relatorio_linhas.append(f"{novo_nome} | PROTOCOLO ORFAO")
            lista_docs.append({
                "ordem": ordem, "nome": novo_nome,
                "original": r.get("nome_original", ""),
                "tipo": "Protocolo Assinatura", "data": None,
            })
            ordem += 1

        relatorio_linhas.append("")
        relatorio_linhas.append("Gerado por: Organizador Juridico AB Group")
        zf.writestr(f"{nome_pasta_zip}/_relatorio.txt", "\n".join(relatorio_linhas))

    zip_buffer.seek(0)
    shared_zip = SHARED_ZIP_DIR / f"{nome_pasta}.zip"
    with open(shared_zip, "wb") as f:
        f.write(zip_buffer.getvalue())

    return {
        "nome_pasta": nome_pasta,
        "lista_docs": lista_docs,
        "total": len(docs_ativos),
        "com_data": sum(1 for r in docs_ativos if r.get("data")),
        "sem_data": sum(1 for r in docs_ativos if not r.get("data")),
        "avisos": avisos_relatorio,
    }


# ============ DASHBOARD ADMIN (#5.4) ============

def _ler_planilha_csv():
    """Le a planilha de logs publica e retorna lista de dicts.
    Cabecalho esperado: timestamp, usuario, cliente, tipo_processo, total_documentos, status
    """
    try:
        import urllib.request
        import csv
        import io as iolib
        req = urllib.request.Request(SPREADSHEET_CSV_URL, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read().decode("utf-8")
        leitor = csv.DictReader(iolib.StringIO(raw))
        return list(leitor)
    except Exception as e:
        return [{"_erro": str(e)}]


def _calcular_metricas(linhas):
    """Calcula metricas para o dashboard a partir das linhas da planilha."""
    if linhas and "_erro" in linhas[0]:
        return {"erro": linhas[0]["_erro"]}

    from collections import Counter
    from datetime import timedelta

    total = len(linhas)
    sucessos = sum(1 for l in linhas if (l.get("status") or "").lower() == "sucesso")
    erros = total - sucessos

    # Mes atual
    hoje = datetime.now()
    mes_atual_str = hoje.strftime("%Y-%m")
    mes_anterior_str = (hoje.replace(day=1) - timedelta(days=1)).strftime("%Y-%m")

    casos_mes_atual = sum(
        1 for l in linhas
        if (l.get("timestamp") or "").startswith(mes_atual_str)
        and (l.get("status") or "").lower() == "sucesso"
    )
    casos_mes_anterior = sum(
        1 for l in linhas
        if (l.get("timestamp") or "").startswith(mes_anterior_str)
        and (l.get("status") or "").lower() == "sucesso"
    )

    # Por usuario
    por_usuario = Counter()
    for l in linhas:
        if (l.get("status") or "").lower() == "sucesso":
            por_usuario[l.get("usuario") or "desconhecido"] += 1

    # Por tipo de processo
    por_tipo = Counter()
    for l in linhas:
        if (l.get("status") or "").lower() == "sucesso":
            por_tipo[l.get("tipo_processo") or "desconhecido"] += 1

    # Erros recentes (ultimos 20)
    erros_lista = [
        l for l in linhas
        if (l.get("status") or "").startswith("ERRO[") or (l.get("status") or "").startswith("erro:")
    ]
    erros_lista.sort(key=lambda l: l.get("timestamp") or "", reverse=True)
    erros_recentes = erros_lista[:20]

    # Total de documentos processados (soma do campo)
    total_docs = 0
    for l in linhas:
        if (l.get("status") or "").lower() == "sucesso":
            try:
                total_docs += int(l.get("total_documentos") or 0)
            except (ValueError, TypeError):
                pass

    # Custo estimado: ~R$ 0.05 por caso (Haiku + dupla checagem)
    custo_estimado = round(sucessos * 0.05, 2)
    custo_mes_atual = round(casos_mes_atual * 0.05, 2)

    return {
        "total_casos": total,
        "sucessos": sucessos,
        "erros": erros,
        "taxa_sucesso": round(sucessos / total * 100, 1) if total else 0,
        "casos_mes_atual": casos_mes_atual,
        "casos_mes_anterior": casos_mes_anterior,
        "por_usuario": por_usuario.most_common(),
        "por_tipo": por_tipo.most_common(),
        "erros_recentes": erros_recentes,
        "total_documentos": total_docs,
        "custo_estimado_total": custo_estimado,
        "custo_estimado_mes": custo_mes_atual,
        "atualizado_em": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def _check_auth(auth):
    """HTTP Basic Auth: aceita qualquer username, senha em ADMIN_PASSWORD."""
    return auth and auth.password == ADMIN_PASSWORD


def _auth_required():
    return Response(
        "Acesso restrito. Login necessario.\n",
        401,
        {"WWW-Authenticate": 'Basic realm="Admin Dashboard"'},
    )


@app.route("/admin")
def admin_dashboard():
    auth = request.authorization
    if not _check_auth(auth):
        return _auth_required()

    linhas = _ler_planilha_csv()
    metricas = _calcular_metricas(linhas)
    return render_template("admin.html", m=metricas)


@app.route("/admin/api/metricas")
def admin_api_metricas():
    """Endpoint JSON pras metricas (usado pelo dashboard pra refresh sem recarregar a pagina)."""
    auth = request.authorization
    if not _check_auth(auth):
        return _auth_required()
    linhas = _ler_planilha_csv()
    return jsonify(_calcular_metricas(linhas))


@app.route("/admin/audit")
def admin_audit():
    """Auditoria local (SQLite) — sobrevive mesmo se o n8n estiver quebrado.
    Mostra todos os eventos (INICIO, SUCESSO, ERRO) dos ultimos N dias.
    Filtros: ?dias=N (default 7), ?evento=ERRO|INICIO|SUCESSO, ?usuario=X
    """
    auth = request.authorization
    if not _check_auth(auth):
        return _auth_required()

    dias = int(request.args.get("dias", "7"))
    filtro_evento = request.args.get("evento", "").upper().strip()
    filtro_usuario = request.args.get("usuario", "").strip()

    todos = audit_query(dias=dias, limit=2000)
    if filtro_evento:
        todos = [r for r in todos if (r.get("evento") or "").upper() == filtro_evento]
    if filtro_usuario:
        todos = [r for r in todos if (r.get("usuario") or "").lower() == filtro_usuario.lower()]

    # Resumo rapido
    from collections import Counter
    eventos_count = Counter(r.get("evento") for r in todos)
    erros_por_tipo = Counter(r.get("tipo_erro") for r in todos if r.get("evento") == "ERRO" and r.get("tipo_erro"))
    usuarios_count = Counter(r.get("usuario") for r in todos if r.get("usuario"))

    return render_template(
        "admin_audit.html",
        registros=todos,
        eventos_count=dict(eventos_count),
        erros_por_tipo=erros_por_tipo.most_common(),
        usuarios_count=usuarios_count.most_common(),
        dias=dias,
        filtro_evento=filtro_evento,
        filtro_usuario=filtro_usuario,
        total=len(todos),
    )


@app.route("/admin/audit.json")
def admin_audit_json():
    """Versao JSON do audit pra integracoes."""
    auth = request.authorization
    if not _check_auth(auth):
        return _auth_required()
    dias = int(request.args.get("dias", "7"))
    return jsonify(audit_query(dias=dias, limit=5000))


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
